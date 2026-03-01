"""Server-side DOCX generation using python-docx.

Converts a Word document spec (sections with heading/paragraph/list/table)
into a binary DOCX file.
"""

import io
import logging
import re
from typing import Any

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


def _strip_html(html: str) -> str:
    """Strip HTML tags and decode basic entities, returning plain text."""
    text = re.sub(r"<br\s*/?>", "\n", html)
    text = re.sub(r"<[^>]+>", "", text)
    text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    text = text.replace("&quot;", '"').replace("&#39;", "'").replace("&nbsp;", " ")
    return text


def _add_rich_text(paragraph, html: str) -> None:
    """Parse simple HTML (bold, italic, underline) into runs on *paragraph*.

    Handles <b>/<strong>, <i>/<em>, <u> tags and plain text segments.
    Nested tags are not fully supported — keeps it simple.
    """
    # Split on tags, keeping the tags
    parts = re.split(r"(</?(?:b|strong|i|em|u|s|br)(?:\s[^>]*)?>)", html, flags=re.I)
    bold = False
    italic = False
    underline = False
    for part in parts:
        lower = part.lower().strip()
        if lower in ("<b>", "<strong>"):
            bold = True
            continue
        if lower in ("</b>", "</strong>"):
            bold = False
            continue
        if lower in ("<i>", "<em>"):
            italic = True
            continue
        if lower in ("</i>", "</em>"):
            italic = False
            continue
        if lower == "<u>":
            underline = True
            continue
        if lower == "</u>":
            underline = False
            continue
        if lower in ("<br>", "<br/>", "<br />"):
            paragraph.add_run("\n")
            continue
        if lower.startswith("<") and lower.endswith(">"):
            continue  # skip other tags
        # Plain text — decode entities
        text = part.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
        text = text.replace("&quot;", '"').replace("&#39;", "'").replace("&nbsp;", " ")
        if not text:
            continue
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.underline = underline


def export_docx(spec: dict) -> bytes:
    """Generate DOCX from a Word document spec.

    Args:
        spec: Word document spec with ``sections`` list. Each section has:
            - ``type``: 'heading', 'paragraph', 'list', 'table'
            - ``content``: HTML string (for heading/paragraph)
            - ``level``: int (for heading, 1-6)
            - ``items``: list of strings (for list)
            - ``ordered``: bool (for list)
            - ``headers``: list of strings (for table)
            - ``rows``: list of list of strings (for table)

    Returns:
        DOCX file content as bytes.
    """
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    sections = spec.get("sections", [])
    title = spec.get("title", "")

    if title and not any(
        s.get("type") == "heading" and s.get("level", 1) == 1 for s in sections
    ):
        doc.add_heading(title, level=1)

    for section in sections:
        sec_type = section.get("type", "paragraph")
        content = section.get("content", "")

        if sec_type == "heading":
            level = min(max(section.get("level", 1), 1), 6)
            # DOCX supports heading levels 0-9 (0 = Title)
            docx_level = min(level, 9)
            heading = doc.add_heading(level=docx_level)
            _add_rich_text(heading, content)

        elif sec_type == "paragraph":
            para = doc.add_paragraph()
            _add_rich_text(para, content)

        elif sec_type == "list":
            items = section.get("items") or []
            if not items and content:
                items = content.split("\n")
            ordered = section.get("ordered", False)
            style_name = "List Number" if ordered else "List Bullet"
            for item in items:
                para = doc.add_paragraph(style=style_name)
                _add_rich_text(para, item)

        elif sec_type == "table":
            # Table data can be at section level or nested under content
            tbl = section
            if isinstance(section.get("content"), dict) and ("headers" in section["content"] or "rows" in section["content"]):
                tbl = section["content"]
            headers = tbl.get("headers", [])
            rows = tbl.get("rows", [])
            if not headers and rows:
                cols = len(rows[0]) if rows else 1
            else:
                cols = len(headers)
            total_rows = (1 if headers else 0) + len(rows)
            if total_rows == 0:
                continue

            table = doc.add_table(rows=total_rows, cols=cols)
            table.style = "Table Grid"

            row_idx = 0
            if headers:
                for col_idx, header in enumerate(headers):
                    cell = table.cell(0, col_idx)
                    cell.text = _strip_html(str(header))
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                row_idx = 1

            for data_row in rows:
                for col_idx, cell_val in enumerate(data_row):
                    if col_idx < cols:
                        cell = table.cell(row_idx, col_idx)
                        cell.text = _strip_html(str(cell_val)) if cell_val is not None else ""
                row_idx += 1

        elif sec_type == "image":
            # Inline image (e.g. chart rendered to PNG on the client)
            img_data = section.get("data", "")
            if img_data.startswith("data:"):
                import base64 as _b64
                # Strip data URL prefix: data:image/png;base64,AAAA...
                b64_part = img_data.split(",", 1)[1] if "," in img_data else ""
                if b64_part:
                    img_bytes = _b64.b64decode(b64_part)
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(6.0))

        else:
            # Unknown section type — treat as paragraph
            if content:
                para = doc.add_paragraph()
                _add_rich_text(para, content)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
