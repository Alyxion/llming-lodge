"""Server integration helpers (framework-agnostic).

Host apps call these to register static assets and API routes.

**No NiceGUI dependency.**  Auth cookies are set either via HTTP
``Set-Cookie`` headers or via the host app's own mechanism.
"""

import os
import logging
import threading
from typing import Optional

# Auth — delegated to llming-com shared library
from llming_com.auth import (  # noqa: F401 — re-exported for backward compat
    AUTH_COOKIE_NAME,
    SESSION_COOKIE_NAME,
    IDENTITY_COOKIE_NAME,
    sign_auth_token,
    verify_auth_cookie,
    sign_identity_token,
    verify_identity_cookie,
    make_auth_cookie_value,
)

logger = logging.getLogger(__name__)

# API path constants
API_PREFIX = "/api/llming"
STATIC_PREFIX = "/llming-static"
API_UPLOAD = f"{API_PREFIX}/upload"
API_IMAGE_PASTE = f"{API_PREFIX}/image-paste"
API_LOAD_CONVERSATION = f"{API_PREFIX}/load-conversation"
ASSET_URL_PREFIX = f"{API_PREFIX}/asset"


class SessionDataStore:
    """Thread-safe store for session-scoped runtime data.

    Manages:
    - **assets**: binary blobs served via HTTP (e.g. contact photos from MCP tools)
    - **pasted_images**: base64 images posted by JS before the WS message references them
    - **pending_loads**: conversation data posted by JS for server-side hydration
    """

    _lock = threading.Lock()
    _assets: dict[str, tuple[bytes, str]] = {}
    _pasted_images: dict[str, list[str]] = {}
    _pending_loads: dict[str, dict] = {}

    # ── Assets ────────────────────────────────────────────────

    @classmethod
    def put_asset(cls, path: str, data: bytes, content_type: str) -> None:
        with cls._lock:
            cls._assets[path] = (data, content_type)

    @classmethod
    def get_asset(cls, path: str) -> Optional[tuple[bytes, str]]:
        with cls._lock:
            return cls._assets.get(path)

    # ── Pasted images ─────────────────────────────────────────

    @classmethod
    def set_pasted_images(cls, session_id: str, images: list[str]) -> None:
        with cls._lock:
            cls._pasted_images[session_id] = images

    @classmethod
    def pop_pasted_images(cls, session_id: str) -> list[str]:
        with cls._lock:
            return cls._pasted_images.pop(session_id, [])

    @classmethod
    def clear_pasted_images(cls, session_id: str) -> None:
        with cls._lock:
            cls._pasted_images.pop(session_id, None)

    # ── Pending conversation loads ────────────────────────────

    @classmethod
    def set_pending_load(cls, session_id: str, data: dict) -> None:
        with cls._lock:
            cls._pending_loads[session_id] = data

    @classmethod
    def pop_pending_load(cls, session_id: str) -> Optional[dict]:
        with cls._lock:
            return cls._pending_loads.pop(session_id, None)


def build_theme_css(accent: str) -> str:
    """Generate CSS overrides for the chat theme accent color.

    Args:
        accent: Hex color string, e.g. '#003D8F'

    Returns:
        CSS string that overrides the chat accent variables for both
        light and dark mode.
    """
    hex_color = accent.lstrip('#')
    r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)

    # Darken by ~25% for hover state
    hover_r = max(0, int(r * 0.75))
    hover_g = max(0, int(g * 0.75))
    hover_b = max(0, int(b * 0.75))
    hover_hex = f"#{hover_r:02x}{hover_g:02x}{hover_b:02x}"

    return (
        f":root, #chat-app.cv2-dark {{\n"
        f"  --chat-accent: {accent};\n"
        f"  --chat-accent-rgb: {r}, {g}, {b};\n"
        f"  --chat-accent-hover: {hover_hex};\n"
        f"  --chat-accent-light: rgba({r}, {g}, {b}, 0.12);\n"
        f"}}"
    )


def get_static_path() -> str:
    """Return absolute path to the static assets directory."""
    return os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static')


def get_chat_static_path() -> str:
    """Return absolute path to the chat static assets directory."""
    return os.path.join(get_static_path(), 'chat')


def _xlsx_preview(data: bytes) -> dict:
    """Convert XLSX bytes to structured table preview data."""
    import io
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    sheets = []
    for ws in wb.worksheets[:5]:
        cols, rows = [], []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                cols = [str(c) if c is not None else f'Col {j+1}' for j, c in enumerate(row)]
            else:
                rows.append([str(c) if c is not None else '' for c in row])
            if i >= 100:
                break
        if not cols:
            continue
        sheets.append({
            "name": ws.title, "columns": cols, "rows": rows,
            "total_rows": ws.max_row or 0,
        })
    wb.close()
    return {"type": "table", "sheets": sheets}


def _docx_preview(data: bytes) -> dict:
    """Convert DOCX bytes to HTML preview."""
    import io
    from docx import Document
    from html import escape
    doc = Document(io.BytesIO(data))
    parts = []
    for para in doc.paragraphs[:300]:
        text = para.text
        if not text.strip():
            parts.append('<br>')
            continue
        sn = para.style.name if para.style else ''
        if 'Heading 1' in sn:
            parts.append(f'<h1>{escape(text)}</h1>')
        elif 'Heading 2' in sn:
            parts.append(f'<h2>{escape(text)}</h2>')
        elif 'Heading 3' in sn:
            parts.append(f'<h3>{escape(text)}</h3>')
        else:
            runs_html = ''
            for run in para.runs:
                t = escape(run.text)
                if run.bold:
                    t = f'<b>{t}</b>'
                if run.italic:
                    t = f'<i>{t}</i>'
                if run.underline:
                    t = f'<u>{t}</u>'
                runs_html += t
            parts.append(f'<p>{runs_html or escape(text)}</p>')
    for table in doc.tables[:10]:
        parts.append('<table><thead>')
        for i, row in enumerate(table.rows[:50]):
            if i == 0:
                parts.append('<tr>')
                for cell in row.cells:
                    parts.append(f'<th>{escape(cell.text)}</th>')
                parts.append('</tr></thead><tbody>')
            else:
                parts.append('<tr>')
                for cell in row.cells:
                    parts.append(f'<td>{escape(cell.text)}</td>')
                parts.append('</tr>')
        parts.append('</tbody></table>')
    return {"type": "html", "html": '\n'.join(parts)}


def build_http_router():
    """Build the FastAPI APIRouter with file upload, image paste, and asset endpoints.

    This contains the HTTP routes that were previously in ui/chat.py._build_router().
    The host app should include this router once at startup.
    """
    from fastapi import APIRouter, UploadFile, File, Header, HTTPException, Request
    from fastapi.responses import Response
    from pydantic import BaseModel

    from fastapi.responses import JSONResponse
    from llming_lodge.documents import UploadManager
    from llming_lodge.api.chat_session_api import SessionRegistry, _handle_client_message
    import llming_lodge.api.chat_session_api as _chat_api

    router = APIRouter(prefix=API_PREFIX)

    # ---- Rich MCP render spec fetch ----

    @router.get("/rich-render/{render_id}")
    async def get_rich_render(render_id: str):
        """Fetch a stored rich MCP render spec by UUID."""
        coll = _chat_api._rich_mcp_coll_cache
        if coll is None:
            # Try to initialize from any active session
            reg = SessionRegistry.get()
            for entry in (reg._sessions.values() if hasattr(reg, '_sessions') else []):
                coll = _chat_api._get_rich_mcp_coll(entry.controller)
                if coll:
                    break
        if coll is None:
            return JSONResponse({"error": "not_available"}, status_code=503)
        doc = await coll.find_one({"_id": render_id})
        if not doc:
            return JSONResponse({"error": "not_found"}, status_code=404)
        render = doc.get("render", {})
        render["version"] = doc.get("version", "1.0")
        return JSONResponse(render)

    # ---- File uploads ----

    # MIME types that support text extraction fallback for oversized files
    _EXTRACTABLE_MIMES = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }
    _MAX_EXTRACTABLE_SIZE = 50 * 1024 * 1024  # 50 MB hard ceiling for extraction

    @router.post("/upload/{session_id}")
    async def upload_files(
        session_id: str,
        files: list[UploadFile] = File(...),
        x_user_id: str = Header(...),
    ):
        mgr = UploadManager.get(session_id)
        if not mgr or mgr.user_id != x_user_id:
            raise HTTPException(403, "Invalid session")
        results = []
        for f in files:
            content = await f.read()
            mime = UploadManager._guess_mime(f.filename or "")
            oversized = len(content) > UploadManager.max_file_size

            if oversized:
                # Text extraction fallback for PDFs and DOCX
                if mime in _EXTRACTABLE_MIMES and len(content) <= _MAX_EXTRACTABLE_SIZE:
                    from llming_lodge.documents.extract import extract_text
                    try:
                        extracted = extract_text(content, mime)
                    except Exception as exc:
                        raise HTTPException(413,
                            f"File too large ({len(content) / (1024*1024):.1f} MB) "
                            f"and text extraction failed: {exc}")
                    if not extracted or extracted.startswith("[Error"):
                        raise HTTPException(413,
                            f"File too large ({len(content) / (1024*1024):.1f} MB) "
                            f"and text could not be extracted.")
                    # Store as text-only (no raw_data → saves memory)
                    att = await mgr.store_extracted(
                        f.filename or "document", extracted, mime, x_user_id,
                    )
                    results.append({
                        "name": att.name,
                        "size": att.size,
                        "fileId": att.file_id,
                        "mimeType": att.mime_type,
                        "extractedText": True,
                    })
                    continue
                raise HTTPException(413, f"File too large: {f.filename}")

            try:
                att = await mgr.store_file(f.filename, content, x_user_id)
            except ValueError as e:
                raise HTTPException(413, str(e))
            results.append({
                "name": att.name,
                "size": att.size,
                "fileId": att.file_id,
                "mimeType": att.mime_type,
            })
        return {"files": results}

    # ---- File preview (serves uploaded files for hover preview) ----

    @router.get("/file-preview/{session_id}/{file_id}")
    async def file_preview(session_id: str, file_id: str):
        """Serve an uploaded file by its file_id for hover preview."""
        from starlette.responses import Response

        mgr = UploadManager.get(session_id)
        if not mgr:
            raise HTTPException(404, "Session not found")
        for f in mgr.files:
            if f.file_id == file_id:
                if not f.raw_data:
                    raise HTTPException(404, "File data not available")
                return Response(content=f.raw_data, media_type=f.mime_type)
        raise HTTPException(404, "File not found")

    # ---- File content preview (structured for popover) ----

    @router.get("/file-content/{session_id}/{file_id}")
    async def file_content_preview(session_id: str, file_id: str):
        """Return structured preview data for a file (table, html, text, or URL)."""
        import asyncio as _aio

        mgr = UploadManager.get(session_id)
        if not mgr:
            raise HTTPException(404, "Session not found")
        file_att = None
        for f in mgr.files:
            if f.file_id == file_id:
                file_att = f
                break
        if not file_att or not file_att.raw_data:
            raise HTTPException(404, "File not found")

        ct = file_att.mime_type or ''
        preview_url = f"{API_PREFIX}/file-preview/{session_id}/{file_id}"

        # Images → URL
        if ct.startswith('image/'):
            return {"type": "image", "url": preview_url, "name": file_att.name}

        # PDF → URL (browser renders natively)
        if ct == 'application/pdf':
            return {"type": "pdf", "url": preview_url, "name": file_att.name}

        # XLSX → structured table
        if 'spreadsheet' in ct or 'xlsx' in ct:
            try:
                result = await _aio.to_thread(_xlsx_preview, file_att.raw_data)
                result["name"] = file_att.name
                return result
            except Exception as e:
                logger.warning(f"[FILE-CONTENT] XLSX preview failed: {e}")

        # DOCX → HTML
        if 'wordprocessing' in ct or 'docx' in ct:
            try:
                result = await _aio.to_thread(_docx_preview, file_att.raw_data)
                result["name"] = file_att.name
                return result
            except Exception as e:
                logger.warning(f"[FILE-CONTENT] DOCX preview failed: {e}")

        # Text / CSV / JSON
        if ct.startswith('text/') or ct == 'application/json' or ct == 'application/csv':
            try:
                content = file_att.raw_data.decode(errors='replace')[:50000]
                return {"type": "text", "content": content, "name": file_att.name}
            except Exception:
                pass

        # Fallback
        return {
            "type": "download", "url": preview_url,
            "name": file_att.name, "size": file_att.size, "content_type": ct,
        }

    # ---- Image paste ----

    class ImagePastePayload(BaseModel):
        images: list[str] = []

    @router.post("/image-paste/{session_id}")
    async def image_paste(session_id: str, payload: ImagePastePayload):
        SessionDataStore.set_pasted_images(session_id, payload.images)
        logger.info(f"[IMAGE-PASTE-API] Stored {len(payload.images)} images for session {session_id}")
        return {"ok": True, "count": len(payload.images)}

    # ---- Conversation load ----

    @router.post("/load-conversation/{session_id}")
    async def load_conversation(session_id: str, request: Request):
        """Receive conversation data from JS (fetched from IDB) via HTTP POST."""
        data = await request.json()
        SessionDataStore.set_pending_load(session_id, data)
        logger.info(f"[LOAD-API] Received conversation {session_id} ({len(data.get('messages', []))} msgs)")
        return {"ok": True}

    # ---- Session-scoped assets (photos, etc.) ----

    @router.get("/asset/{path:path}")
    async def serve_asset(path: str):
        """Serve a binary asset (e.g. contact photo)."""
        entry = SessionDataStore.get_asset(path)
        if entry is None:
            return Response(status_code=404, content="Not found")
        data, content_type = entry
        return Response(
            content=data,
            media_type=content_type,
            headers={"Cache-Control": "private, max-age=3600"},
        )

    # ---- PPTX template export ----

    @router.post("/pptx/export/{session_id}")
    async def export_pptx_templated(session_id: str, request: Request):
        """Generate PPTX using a real template file. Returns binary PPTX."""
        import asyncio as _aio

        data = await request.json()
        spec = data.get("spec", {})
        chart_images = data.get("chartImages", {})
        template_name = spec.get("template", "")

        # Look up template: try active session first, then global registry
        registry = SessionRegistry.get()
        template = None
        entry = registry.get_session(session_id)
        if entry:
            doc_manager = getattr(entry, "doc_manager", None)
            if doc_manager and template_name:
                for tpl in getattr(doc_manager, "presentation_templates", []):
                    if getattr(tpl, "name", "") == template_name:
                        template = tpl
                        break

        # Fall back to global template registry (works for restored chats)
        if not template and template_name:
            template = registry.get_template(template_name)

        if not template or not getattr(template, "template_path", ""):
            raise HTTPException(400, f"Template '{template_name}' not found or has no template_path")

        from llming_docs.pptx_exporter import export_pptx as _export_pptx

        template_config = template.model_dump(by_alias=True)
        pptx_bytes = await _aio.to_thread(
            _export_pptx, spec, template.template_path, chart_images, template_config,
        )

        filename = (spec.get("title") or "presentation").replace("/", "_").strip() or "presentation"
        return Response(
            content=pptx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
            headers={"Content-Disposition": f'attachment; filename="{filename}.pptx"'},
        )

    # ---- Word/DOCX export ----

    @router.post("/word/export")
    async def export_word_docx(request: Request):
        """Generate DOCX from a Word document spec. Returns binary DOCX."""
        import asyncio as _aio

        data = await request.json()
        spec = data.get("spec", {})

        from llming_docs.word_exporter import export_docx as _export_docx

        docx_bytes = await _aio.to_thread(_export_docx, spec)

        filename = (spec.get("title") or "document").replace("/", "_").strip() or "document"
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}.docx"'},
        )

    # ---- Unified document export ----

    @router.post("/doc/export/{session_id}")
    async def export_document(session_id: str, request: Request):
        """Unified document export endpoint.

        Accepts either:
          {"doc_id": "...", "format": "docx|pptx|xlsx|csv|html|png", "chart_images": {...}}
        or:
          {"spec": {...}, "type": "table", "format": "xlsx", "chart_images": {...}}
        """
        import asyncio as _aio

        from llming_docs.render import (
            RenderContext, render_to as _render_to, can_render as _can_render,
        )

        data = await request.json()
        target_format = data.get("format", "")
        chart_images = data.get("chart_images") or data.get("chartImages") or {}
        spec = data.get("spec")
        doc_type = data.get("type", "")

        if not spec:
            raise HTTPException(400, "Missing 'spec' in request body")
        if not doc_type:
            raise HTTPException(400, "Missing 'type' in request body")
        if not target_format:
            raise HTTPException(400, "Missing 'format' in request body")

        # Resolve type aliases
        _aliases = {"word": "text_doc", "powerpoint": "presentation", "pptx": "presentation"}
        doc_type = _aliases.get(doc_type, doc_type)

        if not _can_render(doc_type, target_format):
            raise HTTPException(400, f"Cannot render '{doc_type}' to '{target_format}'")

        # Build context
        template_path = None
        template_config = None
        if doc_type == "presentation":
            template_name = spec.get("template", "")
            registry = SessionRegistry.get()
            template = None
            entry = registry.get_session(session_id)
            if entry:
                doc_manager = getattr(entry, "doc_manager", None)
                if doc_manager and template_name:
                    for tpl in getattr(doc_manager, "presentation_templates", []):
                        if getattr(tpl, "name", "") == template_name:
                            template = tpl
                            break
            if not template and template_name:
                template = registry.get_template(template_name)
            if not template or not getattr(template, "template_path", ""):
                raise HTTPException(400, f"Template '{template_name}' not found")
            template_path = template.template_path
            template_config = template.model_dump(by_alias=True)

        ctx = RenderContext(
            chart_images=chart_images,
            template_path=template_path,
            template_config=template_config,
            session_id=session_id,
        )

        result = await _aio.to_thread(_render_to, doc_type, spec, target_format, ctx)

        return Response(
            content=result.data,
            media_type=result.content_type,
            headers={"Content-Disposition": f'attachment; filename="{result.filename}"'},
        )

    # ---- WebSocket echo (debug) ----

    from starlette.websockets import WebSocket as _WS, WebSocketDisconnect as _WSD

    @router.websocket("/ws-echo")
    async def ws_echo(ws: _WS):
        """Minimal echo endpoint for debugging WS connectivity."""
        await ws.accept()
        logger.info("[WS-ECHO] Client connected")
        try:
            while True:
                data = await ws.receive_text()
                await ws.send_text(f"echo: {data}")
        except _WSD:
            logger.info("[WS-ECHO] Client disconnected")

    # ---- WebSocket chat endpoint (Chat) ----

    import json as _json
    import time as _time

    @router.websocket("/ws/{session_id}")
    async def websocket_chat(ws: _WS, session_id: str):
        """WebSocket endpoint for Chat static frontend."""
        logger.info(f"[WS] Connection attempt for session {session_id}")

        registry = SessionRegistry.get()
        entry = registry.get_session(session_id)
        if not entry:
            logger.warning(f"[WS] Session {session_id} not found in registry "
                           f"(active sessions: {registry.active_count})")
            await ws.accept()
            await ws.send_json({
                "type": "error",
                "error_type": "SessionNotFound",
                "message": f"Session {session_id} not found",
            })
            await ws.close(code=4004, reason="Session not found")
            return

        await ws.accept()
        controller = entry.controller
        controller.set_websocket(ws)
        entry.websocket = ws

        # Wire condensation callbacks
        controller._wire_condensation()

        # Send session init
        init_msg = await controller.build_session_init(user_name=entry.user_name)
        await ws.send_json(init_msg)

        # Send initial context info
        await controller._send_context_info()

        logger.info(f"[WS] Client connected to session {session_id}")

        try:
            while True:
                raw = await ws.receive_text()
                try:
                    msg = _json.loads(raw)
                except _json.JSONDecodeError:
                    await ws.send_json({"type": "error", "error_type": "InvalidJSON", "message": "Invalid JSON"})
                    continue

                entry.last_activity = _time.monotonic()
                await _handle_client_message(controller, entry, msg)

        except _WSD:
            logger.info(f"[WS] Client disconnected from session {session_id}")
        except Exception as e:
            logger.error(f"[WS] Error in session {session_id}: {type(e).__name__}: {e}")
        finally:
            controller.set_websocket(None)
            entry.websocket = None
            # Full session cleanup (nudges, MCP, uploads, doc plugins)
            from llming_lodge.chat_page import cleanup_session as _cleanup
            await _cleanup(session_id)

    return router


def get_ws_router():
    """Return the combined FastAPI APIRouter for WebSocket + HTTP endpoints.

    Includes: file upload, image paste, conversation load, assets, WS chat.
    """
    return build_http_router()


def setup_routes(app, *, debug: bool = False, nudge_store=None) -> None:
    """Mount all llming-lodge routes (static files + API) on a Starlette/FastAPI app.

    Framework-agnostic — works with any app that supports ``.mount()``
    and ``.include_router()`` (FastAPI, Starlette).

    Mounts:
    - Static files (chat assets)
    - HTTP + WebSocket API router
    - Debug API (if ``debug=True`` and ``DEBUG_CHAT_API_KEY`` env var is set)
    - Dev dashboard (if ``LLMING_DEV_PASSWORD`` env var is set)
    - Public API — droplets + remote chat (if ``nudge_store`` is provided)
    """
    from starlette.staticfiles import StaticFiles

    app.mount(STATIC_PREFIX, StaticFiles(directory=get_static_path()), name="llming-static")
    app.mount("/chat-static", StaticFiles(directory=get_chat_static_path()), name="chat-static")

    app.include_router(build_http_router())

    if debug:
        try:
            from llming_lodge.api.debug_api import build_debug_router
            app.include_router(build_debug_router(nudge_store=nudge_store))
        except Exception as e:
            logger.warning(f"[CHAT] Failed to load debug API: {e}")

        # Command-based debug API + MCP server (new system, runs alongside legacy)
        try:
            from llming_lodge.api.chat_session_api import SessionRegistry
            import llming_com.base_commands  # noqa: F401 — register base commands
            import llming_lodge.api.commands  # noqa: F401 — register lodge commands

            from llming_com.command_router import build_command_router
            from llming_lodge.api.debug_api import _check_auth
            cmd_router = build_command_router(
                SessionRegistry.get(),
                auth_dependency=_check_auth,
                prefix=f"{API_PREFIX}/debug/v2",
                extras={"nudge_store": nudge_store},
            )
            app.include_router(cmd_router)

            from llming_com.mcp_http_server import mount_mcp_server
            mount_mcp_server(
                app, SessionRegistry.get(),
                prefix=f"{API_PREFIX}/mcp",
                extras={"nudge_store": nudge_store},
            )
        except Exception as e:
            logger.warning(f"[CHAT] Failed to load command API / MCP: {e}")

    # Dev dashboard — requires LLMING_DEV_PASSWORD env var
    if os.environ.get("LLMING_DEV_PASSWORD", ""):
        try:
            from llming_lodge.api.dev_dashboard import build_dev_router, mount_dev_static
            dev_router = build_dev_router()
            if dev_router:
                app.include_router(dev_router)
                mount_dev_static(app)
                logger.info("[DEV] Dev dashboard mounted")
        except Exception as e:
            logger.warning(f"[DEV] Failed to load dev dashboard: {e}")

    # Public API — droplets + remote chat (auth via user API keys)
    if nudge_store:
        try:
            from llming_lodge.api.public_api import build_public_api_router
            pub_router = build_public_api_router(nudge_store)
            app.include_router(pub_router)
            logger.info(f"[PUBLIC_API] Mounted public API with {len(pub_router.routes)} routes")
        except Exception as e:
            logger.warning(f"[PUBLIC_API] Failed to load public API: {e}")
