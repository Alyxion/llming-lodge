"""DOCX text extraction using python-docx, preserving hyperlinks."""
import io
from pathlib import Path
from typing import Union
from xml.etree.ElementTree import Element


_WPFX = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_RPFX = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}"


def _para_text_with_links(paragraph) -> str:
    """Extract paragraph text with hyperlinks inlined as [text](url)."""
    rels = paragraph.part.rels
    pieces: list[str] = []

    for child in paragraph._element:
        if child.tag == f"{_WPFX}r":
            # Normal run — just grab text
            for t in child.iter(f"{_WPFX}t"):
                if t.text:
                    pieces.append(t.text)
        elif child.tag == f"{_WPFX}hyperlink":
            # Hyperlink element — extract text and resolve URL
            r_id = child.get(f"{_RPFX}id")
            anchor = child.get(f"{_WPFX}anchor")
            texts = []
            for t in child.iter(f"{_WPFX}t"):
                if t.text:
                    texts.append(t.text)
            display = "".join(texts)

            url = None
            if r_id and r_id in rels:
                url = rels[r_id].target_ref
            elif anchor:
                url = f"#{anchor}"

            if url and display:
                pieces.append(f"[{display}]({url})")
            elif url:
                pieces.append(url)
            else:
                pieces.append(display)

    return "".join(pieces)


def extract_docx(source: Union[Path, bytes]) -> str:
    from docx import Document

    fp = io.BytesIO(source) if isinstance(source, bytes) else source
    doc = Document(fp)
    parts: list[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        text = _para_text_with_links(p)
        if text.strip():
            parts.append(text)

    # Tables → markdown-style rows (with links preserved)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_parts = []
                for p in cell.paragraphs:
                    t = _para_text_with_links(p)
                    if t.strip():
                        cell_parts.append(t)
                cells.append(" ".join(cell_parts))
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts) if parts else "[No text content found in document]"
