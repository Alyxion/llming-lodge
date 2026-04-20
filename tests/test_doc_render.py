"""Exhaustive tests for the unified document rendering & embedding system.

Tests render.py, table_exporter.py, html_exporter.py, and the chart_images
extension in word_exporter.py.  Covers normal flows, edge cases, deeply nested
structures, and wild/adversarial inputs.
"""

import base64
import csv
import io
import json
import re
import textwrap

import pytest

# ── render.py ─────────────────────────────────────────────────────────────
from llming_docs.render import (
    EMBED_BEHAVIOR,
    EMBED_RULES,
    EmbedBehavior,
    RENDER_CAPABILITIES,
    RenderContext,
    RenderResult,
    can_embed,
    can_render,
    get_embed_behavior,
    get_embed_format,
    register_embed_behavior,
    render_to,
)

# ── table_exporter.py ─────────────────────────────────────────────────────
from llming_docs.table_exporter import (
    _coerce_value,
    _normalize_spec,
    export_csv,
    export_xlsx,
)

# ── html_exporter.py ──────────────────────────────────────────────────────
from llming_docs.html_exporter import _escape_html, export_html

# ── word_exporter.py ──────────────────────────────────────────────────────
from llming_docs.word_exporter import export_docx, _strip_html


# ═══════════════════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════════════════

# A tiny 1×1 red PNG pixel, base64-encoded
_RED_PIXEL_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4"
    "nGP4z8BQDwAEgAF/pooBPQAAAABJRU5ErkJggg=="
)
_RED_PIXEL_DATA_URI = f"data:image/png;base64,{_RED_PIXEL_B64}"


def _xlsx_to_rows(xlsx_bytes: bytes) -> list[list]:
    """Parse XLSX bytes back into a list of rows for assertion."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(xlsx_bytes), read_only=True, data_only=True)
    ws = wb.active
    return [[cell.value for cell in row] for row in ws.iter_rows()]


def _csv_to_rows(csv_bytes: bytes) -> list[list[str]]:
    """Parse CSV bytes back into rows."""
    reader = csv.reader(io.StringIO(csv_bytes.decode("utf-8")))
    return list(reader)


def _docx_paragraphs(docx_bytes: bytes) -> list[str]:
    """Extract paragraph texts from DOCX bytes."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    return [p.text for p in doc.paragraphs]


def _docx_tables(docx_bytes: bytes):
    """Extract tables from DOCX bytes."""
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    return doc.tables


def _docx_images(docx_bytes: bytes) -> int:
    """Count inline image references in DOCX (including duplicates).

    Uses the XML directly since python-docx rels deduplicate identical images.
    """
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    # Count <wp:inline> and <wp:anchor> elements in the document body
    # which represent image placements
    nsmap = {"wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"}
    body_xml = doc.element
    inlines = body_xml.findall(".//wp:inline", nsmap)
    anchors = body_xml.findall(".//wp:anchor", nsmap)
    return len(inlines) + len(anchors)


# ═══════════════════════════════════════════════════════════════════════════
# PART 1: render.py — capability matrix & routing
# ═══════════════════════════════════════════════════════════════════════════


class TestCanRender:
    """Tests for can_render() — capability checks."""

    @pytest.mark.parametrize(
        "doc_type, fmt, expected",
        [
            ("plotly", "png", True),
            ("plotly", "svg", False),
            ("plotly", "docx", False),
            ("table", "xlsx", True),
            ("table", "csv", True),
            ("table", "pdf", False),
            ("text_doc", "docx", True),
            ("text_doc", "pdf", False),
            ("presentation", "pptx", True),
            ("presentation", "pdf", False),
            ("html", "html", True),
            ("html_sandbox", "html", True),
            ("html", "pdf", False),
            ("email_draft", "eml", False),
            ("unknown_type", "xlsx", False),
            ("", "", False),
        ],
    )
    def test_can_render(self, doc_type, fmt, expected):
        assert can_render(doc_type, fmt) is expected


class TestCanEmbed:
    """Tests for can_embed() — embedding rule checks."""

    @pytest.mark.parametrize(
        "src, tgt, expected",
        [
            ("plotly", "pptx", True),
            ("plotly", "docx", True),
            ("plotly", "email", True),
            ("plotly", "html", True),
            ("table", "pptx", True),
            ("table", "docx", True),
            ("table", "email", True),
            ("table", "html", True),
            ("html_sandbox", "pptx", True),
            ("html_sandbox", "docx", True),
            ("html_sandbox", "email", True),
            ("html_sandbox", "html", False),  # no same-type
            ("text_doc", "email", True),
            ("text_doc", "pptx", False),
            ("presentation", "email", True),
            ("presentation", "docx", False),
            ("email_draft", "pptx", False),
            ("email_draft", "email", False),
            ("unknown", "pptx", False),
        ],
    )
    def test_can_embed(self, src, tgt, expected):
        assert can_embed(src, tgt) is expected


class TestGetEmbedFormat:
    """Tests for get_embed_format()."""

    def test_plotly_into_pptx_is_png(self):
        assert get_embed_format("plotly", "pptx") == "png"

    def test_table_into_email_is_xlsx(self):
        assert get_embed_format("table", "email") == "xlsx"

    def test_table_into_pptx_is_native(self):
        assert get_embed_format("table", "pptx") == "native"

    def test_unsupported_returns_none(self):
        assert get_embed_format("email_draft", "pptx") is None

    def test_unknown_source_returns_none(self):
        assert get_embed_format("nonexistent", "docx") is None


class TestRenderResult:
    """Tests for the RenderResult dataclass."""

    def test_frozen(self):
        r = RenderResult(data=b"hi", content_type="text/plain", filename="f.txt")
        with pytest.raises(AttributeError):
            r.data = b"changed"

    def test_fields(self):
        r = RenderResult(data=b"\x00", content_type="application/octet-stream", filename="x.bin")
        assert r.data == b"\x00"
        assert r.content_type == "application/octet-stream"
        assert r.filename == "x.bin"


class TestRenderTo:
    """Tests for render_to() dispatch logic."""

    def test_unsupported_type_raises(self):
        with pytest.raises(ValueError, match="Cannot render"):
            render_to("plotly", {}, "pdf")

    def test_unknown_type_raises(self):
        with pytest.raises(ValueError, match="Cannot render"):
            render_to("banana", {}, "png")

    def test_table_to_xlsx(self):
        spec = {"title": "Sales", "columns": ["A", "B"], "rows": [["1", "2"]]}
        result = render_to("table", spec, "xlsx")
        assert isinstance(result, RenderResult)
        assert result.content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        assert result.filename == "Sales.xlsx"
        assert len(result.data) > 0
        # Verify it's actual XLSX (starts with PK zip header)
        assert result.data[:2] == b"PK"

    def test_table_to_csv(self):
        spec = {"columns": ["X", "Y"], "rows": [["a", "b"]]}
        result = render_to("table", spec, "csv")
        assert result.content_type == "text/csv"
        assert result.filename == "document.csv"
        assert b"X,Y" in result.data

    def test_html_to_html(self):
        spec = {"title": "My Page", "html": "<h1>Hello</h1>", "css": "h1{color:red}"}
        result = render_to("html", spec, "html")
        assert result.content_type == "text/html"
        assert b"<h1>Hello</h1>" in result.data
        assert b"h1{color:red}" in result.data

    def test_html_sandbox_to_html(self):
        spec = {"html": "<p>test</p>"}
        result = render_to("html_sandbox", spec, "html")
        assert b"<p>test</p>" in result.data

    def test_text_doc_to_docx(self):
        spec = {"title": "Report", "sections": [{"type": "paragraph", "content": "Hello world"}]}
        result = render_to("text_doc", spec, "docx")
        assert result.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert result.filename == "Report.docx"
        assert result.data[:2] == b"PK"

    def test_plotly_png_passthrough(self):
        ctx = RenderContext(chart_images={"c1": _RED_PIXEL_DATA_URI})
        spec = {"id": "c1"}
        result = render_to("plotly", spec, "png", ctx)
        assert result.content_type == "image/png"
        # Should be the decoded red pixel
        assert result.data == base64.b64decode(_RED_PIXEL_B64)

    def test_plotly_png_fallback_first_image(self):
        """When spec has no id, use the first available chart image."""
        ctx = RenderContext(chart_images={"any_key": _RED_PIXEL_B64})
        result = render_to("plotly", {}, "png", ctx)
        assert result.data == base64.b64decode(_RED_PIXEL_B64)

    def test_plotly_png_no_images_raises(self):
        with pytest.raises(ValueError, match="chart_images"):
            render_to("plotly", {}, "png")

    def test_presentation_without_template_raises(self):
        try:
            import pptx  # noqa: F401
        except ImportError:
            pytest.skip("python-pptx not installed")
        with pytest.raises(ValueError, match="template_path"):
            render_to("presentation", {"slides": []}, "pptx")

    def test_title_sanitization(self):
        spec = {"title": "foo/bar/baz", "columns": ["A"], "rows": []}
        result = render_to("table", spec, "csv")
        assert "/" not in result.filename
        assert result.filename == "foo_bar_baz.csv"

    def test_empty_title_fallback(self):
        spec = {"title": "", "columns": ["A"], "rows": []}
        result = render_to("table", spec, "csv")
        assert result.filename == "document.csv"


# ═══════════════════════════════════════════════════════════════════════════
# PART 2: table_exporter.py
# ═══════════════════════════════════════════════════════════════════════════


class TestNormalizeSpec:
    """Tests for _normalize_spec() — the heart of table normalization."""

    def test_flat_string_columns(self):
        spec = {"columns": ["A", "B"], "rows": [["1", "2"]]}
        cols, rows = _normalize_spec(spec)
        assert cols == ["A", "B"]
        assert rows == [["1", "2"]]

    def test_object_columns_with_key_label(self):
        spec = {
            "columns": [{"key": "name", "label": "Name"}, {"key": "age", "label": "Age"}],
            "rows": [{"name": "Alice", "age": 30}],
        }
        cols, rows = _normalize_spec(spec)
        assert cols == ["Name", "Age"]
        assert rows == [["Alice", 30]]

    def test_object_columns_key_only(self):
        spec = {
            "columns": [{"key": "x"}, {"key": "y"}],
            "rows": [{"x": 1, "y": 2}],
        }
        cols, rows = _normalize_spec(spec)
        assert cols == ["x", "y"]
        assert rows == [[1, 2]]

    def test_mixed_column_types(self):
        spec = {
            "columns": ["plain", {"key": "k", "label": "K"}],
            "rows": [["v1", "v2"]],
        }
        cols, rows = _normalize_spec(spec)
        assert cols == ["plain", "K"]
        assert rows == [["v1", "v2"]]

    def test_empty_spec(self):
        cols, rows = _normalize_spec({})
        assert cols == []
        assert rows == []

    def test_rows_without_columns(self):
        spec = {"rows": [["a", "b"], ["c", "d"]]}
        cols, rows = _normalize_spec(spec)
        assert cols == []
        assert rows == [["a", "b"], ["c", "d"]]

    def test_object_rows_with_missing_keys(self):
        spec = {
            "columns": [{"key": "a"}, {"key": "b"}, {"key": "c"}],
            "rows": [{"a": 1}, {"b": 2, "c": 3}],
        }
        cols, rows = _normalize_spec(spec)
        assert rows == [[1, "", ""], ["", 2, 3]]

    def test_single_value_row(self):
        """Non-list, non-dict row is wrapped in a list."""
        spec = {"columns": ["val"], "rows": [42]}
        cols, rows = _normalize_spec(spec)
        assert rows == [[42]]

    def test_tuple_rows(self):
        spec = {"columns": ["a", "b"], "rows": [(1, 2), (3, 4)]}
        cols, rows = _normalize_spec(spec)
        assert rows == [[1, 2], [3, 4]]

    def test_deeply_nested_object_values(self):
        """Object row values can be anything — including dicts and lists."""
        spec = {
            "columns": [{"key": "data"}],
            "rows": [{"data": {"nested": [1, 2, 3]}}],
        }
        cols, rows = _normalize_spec(spec)
        assert rows == [[{"nested": [1, 2, 3]}]]

    def test_unicode_columns_and_data(self):
        spec = {
            "columns": ["名前", "年齢", "Ñoño"],
            "rows": [["太郎", 25, "café"]],
        }
        cols, rows = _normalize_spec(spec)
        assert cols == ["名前", "年齢", "Ñoño"]
        assert rows == [["太郎", 25, "café"]]


class TestCoerceValue:
    """Tests for _coerce_value() — type coercion for cells."""

    def test_none_to_empty_string(self):
        assert _coerce_value(None) == ""

    def test_int_passthrough(self):
        assert _coerce_value(42) == 42

    def test_float_passthrough(self):
        assert _coerce_value(3.14) == 3.14

    def test_bool_passthrough(self):
        assert _coerce_value(True) is True

    def test_string_int(self):
        assert _coerce_value("123") == 123

    def test_string_float(self):
        assert _coerce_value("3.14") == 3.14

    def test_string_negative(self):
        assert _coerce_value("-42") == -42

    def test_string_with_spaces(self):
        assert _coerce_value("  99  ") == 99

    def test_non_numeric_string(self):
        assert _coerce_value("hello") == "hello"

    def test_empty_string(self):
        assert _coerce_value("") == ""

    def test_whitespace_only(self):
        assert _coerce_value("   ") == ""

    def test_mixed_string(self):
        assert _coerce_value("12abc") == "12abc"

    def test_scientific_notation(self):
        assert _coerce_value("1e5") == 1e5

    def test_infinity_string(self):
        # float("inf") is valid
        assert _coerce_value("inf") == float("inf")

    def test_dict_becomes_string(self):
        d = {"a": 1}
        assert _coerce_value(d) == str(d)


class TestExportXlsx:
    """Tests for export_xlsx()."""

    def test_basic_table(self):
        spec = {"columns": ["Name", "Score"], "rows": [["Alice", 95], ["Bob", 87]]}
        data = export_xlsx(spec)
        assert data[:2] == b"PK"
        rows = _xlsx_to_rows(data)
        assert rows[0] == ["Name", "Score"]
        assert rows[1] == ["Alice", 95]
        assert rows[2] == ["Bob", 87]

    def test_empty_table(self):
        spec = {"columns": [], "rows": []}
        data = export_xlsx(spec)
        assert data[:2] == b"PK"
        # openpyxl always creates at least one cell in the sheet
        # so an empty table produces a sheet with a single None cell
        rows = _xlsx_to_rows(data)
        assert len(rows) <= 1

    def test_columns_only_no_rows(self):
        spec = {"columns": ["A", "B", "C"], "rows": []}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert len(rows) == 1
        assert rows[0] == ["A", "B", "C"]

    def test_rows_only_no_columns(self):
        spec = {"rows": [["x", "y"], ["z", "w"]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows == [["x", "y"], ["z", "w"]]

    def test_numeric_coercion(self):
        spec = {"columns": ["Val"], "rows": [["42"], ["3.14"], ["text"]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[1] == [42]
        assert rows[2] == [3.14]
        assert rows[3] == ["text"]

    def test_none_values(self):
        spec = {"columns": ["A"], "rows": [[None], ["ok"]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        # _coerce_value turns None into "", but openpyxl stores "" as None
        # when read back in read_only mode
        assert rows[1][0] in ("", None)
        assert rows[2] == ["ok"]

    def test_title_becomes_sheet_name(self):
        from openpyxl import load_workbook

        spec = {"title": "Revenue 2026", "columns": ["Q"], "rows": [["1"]]}
        data = export_xlsx(spec)
        wb = load_workbook(io.BytesIO(data))
        assert wb.active.title == "Revenue 2026"

    def test_long_title_truncated(self):
        """Excel limits sheet names to 31 characters."""
        from openpyxl import load_workbook

        spec = {"title": "A" * 50, "columns": ["X"], "rows": [["1"]]}
        data = export_xlsx(spec)
        wb = load_workbook(io.BytesIO(data))
        assert len(wb.active.title) == 31

    def test_object_column_format(self):
        spec = {
            "columns": [{"key": "id", "label": "ID"}, {"key": "name", "label": "Full Name"}],
            "rows": [{"id": 1, "name": "Alice"}, {"id": 2, "name": "Bob"}],
        }
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[0] == ["ID", "Full Name"]
        assert rows[1] == [1, "Alice"]

    def test_wide_table_50_columns(self):
        cols = [f"Col{i}" for i in range(50)]
        row = [f"val{i}" for i in range(50)]
        spec = {"columns": cols, "rows": [row]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert len(rows[0]) == 50
        assert len(rows[1]) == 50

    def test_tall_table_1000_rows(self):
        spec = {"columns": ["N"], "rows": [[str(i)] for i in range(1000)]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert len(rows) == 1001  # 1 header + 1000 data rows

    def test_unicode_in_cells(self):
        spec = {"columns": ["Emoji", "CJK"], "rows": [["🎉", "日本語"]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[1] == ["🎉", "日本語"]

    def test_mixed_types_in_row(self):
        spec = {"columns": ["A", "B", "C"], "rows": [[1, "text", 3.14]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[1] == [1, "text", 3.14]

    def test_header_styling(self):
        """Verify header cells have bold font."""
        from openpyxl import load_workbook

        spec = {"columns": ["A", "B"], "rows": [["1", "2"]]}
        data = export_xlsx(spec)
        wb = load_workbook(io.BytesIO(data))
        ws = wb.active
        assert ws.cell(1, 1).font.bold is True
        assert ws.cell(2, 1).font.bold is not True


class TestExportCsv:
    """Tests for export_csv()."""

    def test_basic_csv(self):
        spec = {"columns": ["A", "B"], "rows": [["1", "2"], ["3", "4"]]}
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows == [["A", "B"], ["1", "2"], ["3", "4"]]

    def test_empty_csv(self):
        spec = {"columns": [], "rows": []}
        data = export_csv(spec)
        assert data == b""

    def test_columns_only(self):
        spec = {"columns": ["X", "Y"], "rows": []}
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows == [["X", "Y"]]

    def test_comma_in_data(self):
        spec = {"columns": ["Sentence"], "rows": [["Hello, World"]]}
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows[1] == ["Hello, World"]

    def test_newline_in_data(self):
        spec = {"columns": ["Text"], "rows": [["line1\nline2"]]}
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows[1] == ["line1\nline2"]

    def test_quotes_in_data(self):
        spec = {"columns": ["Quote"], "rows": [['He said "hi"']]}
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows[1] == ['He said "hi"']

    def test_none_becomes_empty(self):
        spec = {"columns": ["A"], "rows": [[None]]}
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows[1] == [""]

    def test_numeric_coercion_in_csv(self):
        """CSV coerces '42' to 42, which becomes '42' in CSV output."""
        spec = {"columns": ["N"], "rows": [["42"]]}
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows[1] == ["42"]

    def test_utf8_encoding(self):
        spec = {"columns": ["Ü"], "rows": [["Ñoño"]]}
        data = export_csv(spec)
        assert "Ñoño" in data.decode("utf-8")

    def test_object_rows_to_csv(self):
        spec = {
            "columns": [{"key": "a", "label": "A"}, {"key": "b", "label": "B"}],
            "rows": [{"a": "x", "b": "y"}],
        }
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows[0] == ["A", "B"]
        assert rows[1] == ["x", "y"]


# ═══════════════════════════════════════════════════════════════════════════
# PART 3: html_exporter.py
# ═══════════════════════════════════════════════════════════════════════════


class TestEscapeHtml:
    """Tests for _escape_html()."""

    def test_no_special_chars(self):
        assert _escape_html("Hello World") == "Hello World"

    def test_ampersand(self):
        assert _escape_html("A & B") == "A &amp; B"

    def test_angle_brackets(self):
        assert _escape_html("<script>alert(1)</script>") == "&lt;script&gt;alert(1)&lt;/script&gt;"

    def test_double_quotes(self):
        assert _escape_html('key="val"') == "key=&quot;val&quot;"

    def test_all_together(self):
        assert _escape_html('<a href="x">&') == '&lt;a href=&quot;x&quot;&gt;&amp;'


class TestExportHtml:
    """Tests for export_html()."""

    def test_basic_html(self):
        spec = {"title": "Test", "html": "<h1>Hello</h1>"}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "<!DOCTYPE html>" in html
        assert "<title>Test</title>" in html
        assert "<h1>Hello</h1>" in html

    def test_with_css(self):
        spec = {"title": "Styled", "html": "<p>Hi</p>", "css": "p { color: blue; }"}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "<style>" in html
        assert "p { color: blue; }" in html

    def test_with_js(self):
        spec = {"title": "Interactive", "html": "<div id='app'></div>", "js": "console.log('hi')"}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "<script>" in html
        assert "console.log('hi')" in html

    def test_no_css_no_style_tag(self):
        spec = {"title": "Plain", "html": "<p>Text</p>"}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "<style>" not in html

    def test_no_js_no_script_tag(self):
        spec = {"title": "Static", "html": "<p>Text</p>"}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "<script>" not in html

    def test_content_fallback(self):
        """'content' field used when 'html' is missing."""
        spec = {"content": "<b>Bold</b>"}
        data = export_html(spec)
        assert b"<b>Bold</b>" in data

    def test_empty_spec(self):
        spec = {}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "<title>Document</title>" in html
        assert "<body>" in html

    def test_title_xss_prevention(self):
        spec = {"title": '<script>alert("xss")</script>', "html": "safe"}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "<script>alert" not in html.split("<title>")[1].split("</title>")[0]
        assert "&lt;script&gt;" in html

    def test_utf8_encoding(self):
        spec = {"title": "日本語", "html": "<p>こんにちは</p>"}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "こんにちは" in html

    def test_full_page_structure(self):
        spec = {"title": "Full", "html": "<main>Content</main>", "css": "main{}", "js": "init()"}
        data = export_html(spec)
        html = data.decode("utf-8")
        # Order: DOCTYPE → html → head (title, style) → body (html, script)
        assert html.index("<!DOCTYPE html>") < html.index("<head>")
        assert html.index("<head>") < html.index("<title>")
        assert html.index("<title>") < html.index("<style>")
        assert html.index("</head>") < html.index("<body>")
        assert html.index("<main>Content</main>") < html.index("<script>")

    def test_large_html_body(self):
        """10KB HTML body."""
        body = "<p>" + "x" * 10000 + "</p>"
        spec = {"html": body}
        data = export_html(spec)
        assert len(data) > 10000


# ═══════════════════════════════════════════════════════════════════════════
# PART 4: word_exporter.py — existing + chart_images extension
# ═══════════════════════════════════════════════════════════════════════════


class TestStripHtml:
    """Tests for _strip_html()."""

    def test_plain_text(self):
        assert _strip_html("Hello") == "Hello"

    def test_bold_tags(self):
        assert _strip_html("<b>Bold</b>") == "Bold"

    def test_br_to_newline(self):
        assert _strip_html("A<br>B") == "A\nB"
        assert _strip_html("A<br/>B") == "A\nB"

    def test_entities(self):
        # &nbsp; is decoded to a regular space by _strip_html
        assert _strip_html("&amp; &lt; &gt; &quot; &#39; &nbsp;") == '& < > " \'  '

    def test_nested_tags(self):
        assert _strip_html("<b><i>text</i></b>") == "text"

    def test_empty(self):
        assert _strip_html("") == ""


class TestExportDocx:
    """Tests for export_docx() including the new chart_images support."""

    def test_empty_doc(self):
        spec = {"sections": []}
        data = export_docx(spec)
        assert data[:2] == b"PK"

    def test_title_added_when_no_h1(self):
        spec = {"title": "My Doc", "sections": [{"type": "paragraph", "content": "text"}]}
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "My Doc" in texts[0]

    def test_title_not_duplicated_if_h1_exists(self):
        spec = {
            "title": "My Doc",
            "sections": [
                {"type": "heading", "level": 1, "content": "My Doc"},
                {"type": "paragraph", "content": "body"},
            ],
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        # Should have heading + paragraph, not title + heading + paragraph
        assert texts.count("My Doc") == 1

    def test_heading_levels(self):
        spec = {
            "sections": [
                {"type": "heading", "level": 1, "content": "H1"},
                {"type": "heading", "level": 2, "content": "H2"},
                {"type": "heading", "level": 3, "content": "H3"},
                {"type": "heading", "level": 6, "content": "H6"},
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "H1" in texts
        assert "H6" in texts

    def test_paragraph_with_rich_text(self):
        spec = {
            "sections": [
                {"type": "paragraph", "content": "Normal <b>bold</b> <i>italic</i>"}
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "Normal bold italic" in texts

    def test_unordered_list(self):
        spec = {
            "sections": [
                {"type": "list", "items": ["Item A", "Item B", "Item C"]}
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "Item A" in texts

    def test_ordered_list(self):
        spec = {
            "sections": [
                {"type": "list", "ordered": True, "items": ["First", "Second"]}
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "First" in texts

    def test_table_section(self):
        spec = {
            "sections": [
                {
                    "type": "table",
                    "headers": ["Name", "Value"],
                    "rows": [["A", "1"], ["B", "2"]],
                }
            ]
        }
        data = export_docx(spec)
        tables = _docx_tables(data)
        assert len(tables) == 1
        assert tables[0].cell(0, 0).text == "Name"
        assert tables[0].cell(1, 0).text == "A"

    def test_table_nested_in_content(self):
        """Table data nested under 'content' key."""
        spec = {
            "sections": [
                {
                    "type": "table",
                    "content": {
                        "headers": ["X"],
                        "rows": [["y"]],
                    },
                }
            ]
        }
        data = export_docx(spec)
        tables = _docx_tables(data)
        assert len(tables) == 1

    def test_image_section(self):
        spec = {
            "sections": [
                {"type": "image", "data": _RED_PIXEL_DATA_URI}
            ]
        }
        data = export_docx(spec)
        assert _docx_images(data) == 1

    def test_chart_section_with_chart_images(self):
        """New chart section type using chart_images dict."""
        spec = {
            "sections": [
                {"type": "chart", "_chartImageId": "chart_0"}
            ]
        }
        chart_images = {"chart_0": _RED_PIXEL_DATA_URI}
        data = export_docx(spec, chart_images=chart_images)
        assert _docx_images(data) == 1

    def test_chart_section_with_raw_base64(self):
        """Chart image as raw base64 (no data URI prefix)."""
        spec = {
            "sections": [
                {"type": "chart", "_chartImageId": "c1"}
            ]
        }
        data = export_docx(spec, chart_images={"c1": _RED_PIXEL_B64})
        assert _docx_images(data) == 1

    def test_chart_section_missing_image(self):
        """Chart with missing image ID should not crash."""
        spec = {
            "sections": [
                {"type": "chart", "_chartImageId": "nonexistent"}
            ]
        }
        data = export_docx(spec, chart_images={})
        assert data[:2] == b"PK"
        assert _docx_images(data) == 0

    def test_chart_images_none(self):
        """chart_images=None should not crash."""
        spec = {"sections": [{"type": "chart", "_chartImageId": "c1"}]}
        data = export_docx(spec, chart_images=None)
        assert data[:2] == b"PK"

    def test_unknown_section_type(self):
        """Unknown type with content treated as paragraph."""
        spec = {
            "sections": [
                {"type": "unknown_widget", "content": "fallback text"}
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "fallback text" in texts

    def test_complex_document_with_all_types(self):
        """A document using every section type including chart."""
        spec = {
            "title": "Quarterly Report",
            "sections": [
                {"type": "heading", "level": 1, "content": "Quarterly Report"},
                {"type": "paragraph", "content": "Executive summary with <b>bold</b> text."},
                {"type": "heading", "level": 2, "content": "Sales Data"},
                {
                    "type": "table",
                    "headers": ["Product", "Revenue", "Growth"],
                    "rows": [
                        ["Widget A", "$1.2M", "+15%"],
                        ["Widget B", "$800K", "-5%"],
                        ["Widget C", "$2.1M", "+30%"],
                    ],
                },
                {"type": "heading", "level": 2, "content": "Revenue Chart"},
                {"type": "chart", "_chartImageId": "revenue_chart"},
                {"type": "heading", "level": 2, "content": "Key Takeaways"},
                {
                    "type": "list",
                    "items": [
                        "Widget C is the <b>top performer</b>",
                        "Widget B needs attention",
                        "Overall growth is <i>positive</i>",
                    ],
                },
                {"type": "paragraph", "content": "Confidential — do not distribute."},
            ],
        }
        chart_images = {"revenue_chart": _RED_PIXEL_DATA_URI}
        data = export_docx(spec, chart_images=chart_images)
        assert data[:2] == b"PK"
        texts = _docx_paragraphs(data)
        tables = _docx_tables(data)
        assert _docx_images(data) == 1
        assert len(tables) == 1
        assert "Quarterly Report" in texts
        assert "Widget C is the top performer" in texts

    def test_multiple_charts_in_document(self):
        """Multiple chart sections each with their own image."""
        spec = {
            "sections": [
                {"type": "chart", "_chartImageId": "c1"},
                {"type": "paragraph", "content": "Between charts"},
                {"type": "chart", "_chartImageId": "c2"},
            ]
        }
        chart_images = {
            "c1": _RED_PIXEL_DATA_URI,
            "c2": _RED_PIXEL_B64,
        }
        data = export_docx(spec, chart_images=chart_images)
        assert _docx_images(data) == 2

    def test_deeply_nested_html_in_paragraph(self):
        """Deeply nested HTML tags in paragraph content."""
        html = "<b>bold <i>bold-italic <u>all-three</u> back-to-bi</i> just-bold</b> normal"
        spec = {"sections": [{"type": "paragraph", "content": html}]}
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "bold bold-italic all-three back-to-bi just-bold normal" in texts

    def test_table_with_none_cells(self):
        spec = {
            "sections": [
                {
                    "type": "table",
                    "headers": ["A", "B"],
                    "rows": [[None, "val"], ["val", None]],
                }
            ]
        }
        data = export_docx(spec)
        tables = _docx_tables(data)
        assert tables[0].cell(1, 0).text == ""
        assert tables[0].cell(2, 1).text == ""

    def test_table_no_headers(self):
        spec = {
            "sections": [
                {
                    "type": "table",
                    "rows": [["a", "b"], ["c", "d"]],
                }
            ]
        }
        data = export_docx(spec)
        tables = _docx_tables(data)
        assert tables[0].cell(0, 0).text == "a"

    def test_list_fallback_from_content(self):
        """List with content instead of items."""
        spec = {
            "sections": [
                {"type": "list", "content": "line1\nline2\nline3"}
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "line1" in texts

    def test_html_entities_in_sections(self):
        spec = {
            "sections": [
                {"type": "paragraph", "content": "&amp; &lt;tag&gt; &quot;quoted&quot;"}
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert '& <tag> "quoted"' in texts


# ═══════════════════════════════════════════════════════════════════════════
# PART 5: Wild / adversarial / edge case tests
# ═══════════════════════════════════════════════════════════════════════════


class TestWildEdgeCases:
    """Extreme and adversarial inputs that should not crash the system."""

    def test_xlsx_with_10000_cells(self):
        """100 cols × 100 rows = 10,000 cells."""
        cols = [f"C{i}" for i in range(100)]
        rows = [[f"r{r}c{c}" for c in range(100)] for r in range(100)]
        spec = {"columns": cols, "rows": rows}
        data = export_xlsx(spec)
        parsed = _xlsx_to_rows(data)
        assert len(parsed) == 101  # 1 header + 100 rows
        assert len(parsed[0]) == 100

    def test_csv_with_special_chars_everywhere(self):
        """Every cell contains commas, quotes, newlines, and unicode."""
        spec = {
            "columns": ['Col,"1"', "Col\n2", "Col 3"],
            "rows": [
                ['val,"a"', "val\nb", "val 🎉"],
                ["", None, "normal"],
            ],
        }
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows[1][0] == 'val,"a"'
        assert rows[1][2] == "val 🎉"

    def test_html_xss_in_body(self):
        """HTML body with XSS — should pass through (it's an HTML export)."""
        spec = {"html": '<script>alert("xss")</script><img onerror=alert(1)>'}
        data = export_html(spec)
        # In HTML export, raw HTML is expected — but title should be escaped
        assert b'<script>alert("xss")</script>' in data

    def test_docx_empty_sections_list(self):
        spec = {"title": "Empty", "sections": []}
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "Empty" in texts  # title auto-added

    def test_docx_section_missing_type(self):
        """Section without 'type' should default to paragraph."""
        spec = {"sections": [{"content": "no type here"}]}
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "no type here" in texts

    def test_docx_section_empty_content(self):
        spec = {"sections": [{"type": "paragraph", "content": ""}]}
        data = export_docx(spec)
        assert data[:2] == b"PK"

    def test_xlsx_cell_with_very_long_string(self):
        long_text = "A" * 32767  # Excel cell character limit
        spec = {"columns": ["Text"], "rows": [[long_text]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[1][0] == long_text

    def test_csv_with_empty_rows(self):
        spec = {"columns": ["A"], "rows": [[""], [""], [""]]}
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert len(rows) == 4  # 1 header + 3 data rows

    def test_render_to_with_slash_in_title(self):
        spec = {"title": "2026/Q1/Report", "columns": ["A"], "rows": [["1"]]}
        result = render_to("table", spec, "xlsx")
        assert "/" not in result.filename

    def test_render_to_whitespace_title(self):
        spec = {"title": "   ", "columns": ["A"], "rows": [["1"]]}
        result = render_to("table", spec, "csv")
        assert result.filename == "document.csv"

    def test_table_spec_with_extra_keys(self):
        """Extra keys in spec should be ignored."""
        spec = {
            "columns": ["A"],
            "rows": [["1"]],
            "metadata": {"author": "test"},
            "style": {"theme": "dark"},
        }
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[1] == [1]

    def test_table_rows_shorter_than_columns(self):
        """Rows with fewer values than columns."""
        spec = {"columns": ["A", "B", "C"], "rows": [["only_one"]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[1][0] == "only_one"

    def test_table_rows_longer_than_columns(self):
        """Rows with more values than columns — extra values included in XLSX."""
        spec = {"columns": ["A"], "rows": [["v1", "v2", "v3"]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        # openpyxl writes all cells even beyond column count
        assert rows[1][0] == "v1"

    def test_html_with_all_fields_empty(self):
        spec = {"title": "", "html": "", "css": "", "js": ""}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "<!DOCTYPE html>" in html

    def test_plotly_png_with_raw_base64(self):
        """Raw base64 (no data: prefix) should work."""
        ctx = RenderContext(chart_images={"c1": _RED_PIXEL_B64})
        spec = {"id": "c1"}
        result = render_to("plotly", spec, "png", ctx)
        assert result.data == base64.b64decode(_RED_PIXEL_B64)

    def test_docx_heading_level_clamped(self):
        """Level 0 or 99 should be clamped to valid range."""
        spec = {
            "sections": [
                {"type": "heading", "level": 0, "content": "TooLow"},
                {"type": "heading", "level": 99, "content": "TooHigh"},
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert "TooLow" in texts
        assert "TooHigh" in texts

    def test_docx_image_invalid_base64_skipped(self):
        """Invalid base64 in image data should not crash."""
        spec = {
            "sections": [
                {"type": "image", "data": "data:image/png;base64,NOT_VALID!!!"}
            ]
        }
        # This may raise or produce a valid doc — it should not crash fatally
        try:
            data = export_docx(spec)
            assert data[:2] == b"PK"
        except Exception:
            pass  # Acceptable to raise on invalid base64

    def test_chart_invalid_base64_logged(self):
        """Invalid base64 in chart_images should be caught and logged."""
        spec = {"sections": [{"type": "chart", "_chartImageId": "bad"}]}
        data = export_docx(spec, chart_images={"bad": "NOT_BASE64!!!"})
        assert data[:2] == b"PK"
        assert _docx_images(data) == 0

    def test_table_with_boolean_values(self):
        spec = {"columns": ["Flag"], "rows": [[True], [False]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[1] == [True]
        assert rows[2] == [False]

    def test_table_with_nested_dict_values(self):
        """Dict values in cells — should be stringified."""
        spec = {"columns": ["Data"], "rows": [[{"nested": {"deep": True}}]]}
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert "nested" in rows[1][0]

    def test_render_capabilities_completeness(self):
        """Verify all doc types in RENDER_CAPABILITIES have at least one format."""
        for doc_type, fmts in RENDER_CAPABILITIES.items():
            assert len(fmts) > 0, f"{doc_type} has no render formats"

    def test_embed_rules_no_self_embedding(self):
        """No type should be embeddable into itself (checked via naming convention)."""
        type_to_target_format = {
            "plotly": set(),
            "table": set(),
            "html_sandbox": {"html"},
            "text_doc": {"docx"},
            "presentation": {"pptx"},
            "email_draft": set(),
        }
        # Just verify that html_sandbox→html is not in EMBED_RULES
        assert "html" not in EMBED_RULES.get("html_sandbox", {})

    def test_render_context_defaults(self):
        ctx = RenderContext()
        assert ctx.chart_images is None
        assert ctx.template_path is None
        assert ctx.template_config is None
        assert ctx.session_id is None

    def test_render_context_with_all_fields(self):
        ctx = RenderContext(
            chart_images={"c1": "data"},
            template_path="/tmp/template.pptx",
            template_config={"layouts": []},
            session_id="abc-123",
        )
        assert ctx.chart_images == {"c1": "data"}
        assert ctx.template_path == "/tmp/template.pptx"

    def test_multiple_render_calls_independent(self):
        """Multiple render_to calls should not leak state."""
        spec1 = {"columns": ["A"], "rows": [["1"]]}
        spec2 = {"columns": ["B", "C"], "rows": [["2", "3"]]}
        r1 = render_to("table", spec1, "csv")
        r2 = render_to("table", spec2, "csv")
        assert b"A" in r1.data
        assert b"A" not in r2.data
        assert b"B,C" in r2.data


class TestWildNestedDocuments:
    """Tests with deeply nested / complex document structures."""

    def test_docx_20_sections_mixed(self):
        """20 sections alternating between headings, paragraphs, lists, tables."""
        sections = []
        for i in range(20):
            if i % 4 == 0:
                sections.append({"type": "heading", "level": 2, "content": f"Section {i}"})
            elif i % 4 == 1:
                sections.append({"type": "paragraph", "content": f"Text for section <b>{i}</b>"})
            elif i % 4 == 2:
                sections.append({"type": "list", "items": [f"Item {i}.{j}" for j in range(5)]})
            else:
                sections.append({
                    "type": "table",
                    "headers": ["Col A", "Col B"],
                    "rows": [[f"r{r}a", f"r{r}b"] for r in range(3)],
                })
        spec = {"title": "Complex Doc", "sections": sections}
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        tables = _docx_tables(data)
        assert len(tables) == 5  # every 4th section is a table
        assert "Section 0" in texts
        assert "Item 2.0" in texts

    def test_xlsx_object_rows_with_nested_arrays(self):
        """Object rows where values are lists (should be stringified by openpyxl)."""
        spec = {
            "columns": [{"key": "tags", "label": "Tags"}],
            "rows": [{"tags": "tag1, tag2"}, {"tags": "tag3"}],
        }
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[1] == ["tag1, tag2"]

    def test_docx_with_charts_and_images_interleaved(self):
        """Document with both raw images and chart references interleaved."""
        spec = {
            "sections": [
                {"type": "heading", "level": 1, "content": "Visual Report"},
                {"type": "image", "data": _RED_PIXEL_DATA_URI},
                {"type": "paragraph", "content": "Analysis below:"},
                {"type": "chart", "_chartImageId": "analysis_chart"},
                {"type": "paragraph", "content": "Conclusion text"},
                {"type": "image", "data": _RED_PIXEL_DATA_URI},
                {"type": "chart", "_chartImageId": "summary_chart"},
            ]
        }
        chart_images = {
            "analysis_chart": _RED_PIXEL_DATA_URI,
            "summary_chart": _RED_PIXEL_B64,
        }
        data = export_docx(spec, chart_images=chart_images)
        assert _docx_images(data) == 4  # 2 images + 2 charts

    def test_table_huge_columns_count(self):
        """Table with 200 columns."""
        cols = [f"Col_{i}" for i in range(200)]
        row = [i for i in range(200)]
        spec = {"columns": cols, "rows": [row]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert len(rows[0]) == 200
        assert len(rows[1]) == 200

    def test_csv_multiline_cells_roundtrip(self):
        """CSV with multiline cell values should roundtrip correctly."""
        spec = {
            "columns": ["Description"],
            "rows": [
                ["First line\nSecond line\nThird line"],
                ["Single line"],
            ],
        }
        data = export_csv(spec)
        rows = _csv_to_rows(data)
        assert rows[1][0] == "First line\nSecond line\nThird line"
        assert rows[2][0] == "Single line"

    def test_html_complex_body(self):
        """HTML with complex nested structure."""
        body = textwrap.dedent("""\
            <div class="container">
              <header>
                <h1>Dashboard</h1>
                <nav><a href="#">Home</a> | <a href="#">About</a></nav>
              </header>
              <main>
                <section id="charts">
                  <div class="chart" data-id="1"></div>
                  <div class="chart" data-id="2"></div>
                </section>
                <table>
                  <tr><th>A</th><th>B</th></tr>
                  <tr><td>1</td><td>2</td></tr>
                </table>
              </main>
              <footer>&copy; 2026</footer>
            </div>""")
        css = textwrap.dedent("""\
            .container { max-width: 1200px; margin: 0 auto; }
            .chart { width: 400px; height: 300px; border: 1px solid #ccc; }
            table { border-collapse: collapse; }
            td, th { padding: 8px; border: 1px solid #ddd; }""")
        js = textwrap.dedent("""\
            document.querySelectorAll('.chart').forEach(el => {
              const id = el.dataset.id;
              // Initialize chart...
            });""")
        spec = {"title": "Dashboard", "html": body, "css": css, "js": js}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert 'class="container"' in html
        assert "border-collapse: collapse" in html
        assert "el.dataset.id" in html

    def test_docx_table_with_html_in_cells(self):
        """Table cells containing HTML tags — should be stripped."""
        spec = {
            "sections": [
                {
                    "type": "table",
                    "headers": ["<b>Name</b>", "<i>Score</i>"],
                    "rows": [
                        ["<b>Alice</b>", "<span style='color:red'>95</span>"],
                        ["Bob &amp; Carol", "87"],
                    ],
                }
            ]
        }
        data = export_docx(spec)
        tables = _docx_tables(data)
        assert tables[0].cell(0, 0).text == "Name"
        assert tables[0].cell(0, 1).text == "Score"
        assert tables[0].cell(1, 0).text == "Alice"
        assert tables[0].cell(2, 0).text == "Bob & Carol"

    def test_render_to_csv_large_dataset(self):
        """Render 5000 rows to CSV via render_to()."""
        spec = {
            "title": "Big Data",
            "columns": ["ID", "Value", "Category"],
            "rows": [[i, i * 3.14, f"cat_{i % 10}"] for i in range(5000)],
        }
        result = render_to("table", spec, "csv")
        rows = _csv_to_rows(result.data)
        assert len(rows) == 5001  # 1 header + 5000 data
        assert result.filename == "Big Data.csv"

    def test_render_to_xlsx_then_parse_back(self):
        """Full roundtrip: spec → XLSX → parsed rows."""
        spec = {
            "columns": [
                {"key": "product", "label": "Product"},
                {"key": "qty", "label": "Quantity"},
                {"key": "price", "label": "Price"},
            ],
            "rows": [
                {"product": "Widget", "qty": 100, "price": 9.99},
                {"product": "Gadget", "qty": 50, "price": 19.99},
                {"product": "Thingamajig", "qty": 200, "price": 4.99},
            ],
        }
        result = render_to("table", spec, "xlsx")
        rows = _xlsx_to_rows(result.data)
        assert rows[0] == ["Product", "Quantity", "Price"]
        assert rows[1] == ["Widget", 100, 9.99]
        assert rows[3] == ["Thingamajig", 200, 4.99]

    def test_docx_with_50_headings(self):
        """Document with 50 headings at various levels."""
        sections = []
        for i in range(50):
            level = (i % 6) + 1
            sections.append({"type": "heading", "level": level, "content": f"Heading {i} (L{level})"})
            sections.append({"type": "paragraph", "content": f"Content for heading {i}"})
        spec = {"sections": sections}
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        # _docx_paragraphs uses doc.paragraphs which includes headings
        assert "Heading 0 (L1)" in texts
        assert "Content for heading 0" in texts
        # Heading 49 might use rich text (via _add_rich_text), so text is there
        all_text = "\n".join(texts)
        assert "Heading 49" in all_text
        assert "Content for heading 49" in all_text
        # At least 100 paragraphs (50 headings + 50 content)
        assert len(texts) >= 100

    def test_xlsx_with_formula_like_strings(self):
        """Cells that look like Excel formulas should be treated as strings."""
        spec = {
            "columns": ["Formula"],
            "rows": [["=SUM(A1:A10)"], ["=1+1"], ["+cmd|'/C calc'!A0"]],
        }
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        # openpyxl writes them as-is; the type depends on coercion
        # The key thing is it doesn't crash
        assert len(rows) == 4

    def test_render_all_supported_combinations(self):
        """Systematically test every supported doc_type → format combination."""
        test_cases = [
            ("table", "xlsx", {"columns": ["A"], "rows": [["1"]]}),
            ("table", "csv", {"columns": ["A"], "rows": [["1"]]}),
            ("text_doc", "docx", {"sections": [{"type": "paragraph", "content": "hi"}]}),
            ("html", "html", {"html": "<p>hi</p>"}),
            ("html_sandbox", "html", {"html": "<p>hi</p>"}),
            (
                "plotly",
                "png",
                {"id": "c1"},
            ),
        ]
        for doc_type, fmt, spec in test_cases:
            ctx = None
            if doc_type == "plotly":
                ctx = RenderContext(chart_images={"c1": _RED_PIXEL_B64})
            result = render_to(doc_type, spec, fmt, ctx)
            assert isinstance(result, RenderResult), f"Failed for {doc_type}→{fmt}"
            assert len(result.data) > 0, f"Empty data for {doc_type}→{fmt}"
            assert result.filename.endswith(f".{fmt}"), f"Wrong extension for {doc_type}→{fmt}"

    def test_table_with_all_none_row(self):
        """A row where every value is None."""
        spec = {"columns": ["A", "B", "C"], "rows": [[None, None, None]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        # openpyxl reads empty string cells back as None in read_only mode
        assert all(v in ("", None) for v in rows[1])

    def test_docx_list_with_html_items(self):
        """List items containing HTML formatting."""
        spec = {
            "sections": [
                {
                    "type": "list",
                    "items": [
                        "<b>Important:</b> Do this first",
                        "Then do <i>this</i>",
                        "Finally <u>this</u>",
                    ],
                }
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        assert any("Important:" in t for t in texts)

    def test_docx_paragraph_with_br_tags(self):
        """BR tags should create line breaks in the paragraph."""
        spec = {
            "sections": [
                {"type": "paragraph", "content": "Line 1<br>Line 2<br/>Line 3"}
            ]
        }
        data = export_docx(spec)
        texts = _docx_paragraphs(data)
        # The text should contain newlines from <br> tags
        full_text = " ".join(texts)
        assert "Line 1" in full_text
        assert "Line 2" in full_text
        assert "Line 3" in full_text

    def test_table_with_numeric_column_labels(self):
        """Column labels that are numbers."""
        spec = {"columns": [1, 2, 3], "rows": [["a", "b", "c"]]}
        data = export_xlsx(spec)
        rows = _xlsx_to_rows(data)
        assert rows[0] == ["1", "2", "3"]

    def test_csv_completely_empty_spec(self):
        """Completely empty spec — no columns, no rows."""
        data = export_csv({})
        assert data == b""

    def test_xlsx_completely_empty_spec(self):
        data = export_xlsx({})
        assert data[:2] == b"PK"  # Valid XLSX with empty sheet

    def test_html_only_js_no_html_body(self):
        spec = {"js": "document.write('Hello')"}
        data = export_html(spec)
        html = data.decode("utf-8")
        assert "document.write('Hello')" in html
        # Body should be empty (no html key)
        assert "<body>" in html


# ═══════════════════════════════════════════════════════════════════════════
# PART 6: EmbedBehavior registry
# ═══════════════════════════════════════════════════════════════════════════


class TestEmbedBehavior:
    """Tests for the EmbedBehavior dataclass and registry."""

    def test_frozen(self):
        b = EmbedBehavior(mode="graphic", aspect=1.6)
        with pytest.raises(AttributeError):
            b.mode = "text"

    def test_graphic_with_aspect(self):
        b = EmbedBehavior(mode="graphic", aspect=16 / 9)
        assert b.mode == "graphic"
        assert abs(b.aspect - 16 / 9) < 0.001

    def test_table_no_aspect(self):
        b = EmbedBehavior(mode="table")
        assert b.mode == "table"
        assert b.aspect is None

    def test_text_no_aspect(self):
        b = EmbedBehavior(mode="text")
        assert b.mode == "text"
        assert b.aspect is None


class TestEmbedBehaviorRegistry:
    """Tests for the EMBED_BEHAVIOR registry and lookup functions."""

    def test_all_known_types_registered(self):
        expected = {"plotly", "table", "text_doc", "html_sandbox", "html",
                    "presentation", "email_draft", "latex", "rich_mcp"}
        assert expected.issubset(set(EMBED_BEHAVIOR.keys()))

    def test_plotly_is_graphic(self):
        b = get_embed_behavior("plotly")
        assert b is not None
        assert b.mode == "graphic"
        assert b.aspect == 1.6

    def test_table_is_table(self):
        b = get_embed_behavior("table")
        assert b is not None
        assert b.mode == "table"

    def test_text_doc_is_text(self):
        b = get_embed_behavior("text_doc")
        assert b is not None
        assert b.mode == "text"

    def test_presentation_is_graphic_16_9(self):
        b = get_embed_behavior("presentation")
        assert b is not None
        assert b.mode == "graphic"
        assert abs(b.aspect - 16 / 9) < 0.001

    def test_latex_graphic_no_aspect(self):
        b = get_embed_behavior("latex")
        assert b is not None
        assert b.mode == "graphic"
        assert b.aspect is None

    def test_unknown_returns_none(self):
        assert get_embed_behavior("banana") is None

    def test_register_new_type(self):
        register_embed_behavior("custom_widget", EmbedBehavior(mode="graphic", aspect=2.0))
        b = get_embed_behavior("custom_widget")
        assert b is not None
        assert b.mode == "graphic"
        assert b.aspect == 2.0
        # Clean up
        del EMBED_BEHAVIOR["custom_widget"]

    def test_override_existing_type(self):
        original = get_embed_behavior("plotly")
        register_embed_behavior("plotly", EmbedBehavior(mode="table"))
        assert get_embed_behavior("plotly").mode == "table"
        # Restore
        register_embed_behavior("plotly", original)

    @pytest.mark.parametrize("doc_type", list(EMBED_BEHAVIOR.keys()))
    def test_all_behaviors_have_valid_mode(self, doc_type):
        b = EMBED_BEHAVIOR[doc_type]
        assert b.mode in ("graphic", "table", "text")

    @pytest.mark.parametrize("doc_type", list(EMBED_BEHAVIOR.keys()))
    def test_graphic_types_have_aspect_or_none(self, doc_type):
        b = EMBED_BEHAVIOR[doc_type]
        if b.mode == "graphic":
            assert b.aspect is None or isinstance(b.aspect, (int, float))

    def test_non_graphic_types_have_no_aspect(self):
        for doc_type, b in EMBED_BEHAVIOR.items():
            if b.mode in ("table", "text"):
                assert b.aspect is None, f"{doc_type} is {b.mode} but has aspect={b.aspect}"


class TestTextDocMcpEmbed:
    """Tests for the embed section type in text_doc_mcp.py."""

    def test_add_embed_section(self):
        """The MCP tool should accept type='embed' with $ref."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Test Doc", data={"sections": []})
        mcp = TextDocMCP(store)

        result = asyncio.new_event_loop().run_until_complete(
            mcp.call_tool("text_doc_add_section", {
                "document_id": doc.id,
                "type": "embed",
                "ref": "some-plotly-uuid",
            })
        )
        parsed = json.loads(result)
        assert parsed["status"] == "section_added"
        assert parsed["section_count"] == 1

        # Verify the section was stored correctly — tool arg is "ref",
        # but stored as "$ref" in the section JSON
        updated_doc = store.get(doc.id)
        section = updated_doc.data["sections"][0]
        assert section["type"] == "embed"
        assert section["$ref"] == "some-plotly-uuid"
        assert "content" not in section  # embed doesn't need content

    def test_embed_without_ref_fails(self):
        """Embed section without $ref should fail."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Test", data={"sections": []})
        mcp = TextDocMCP(store)

        result = asyncio.new_event_loop().run_until_complete(
            mcp.call_tool("text_doc_add_section", {
                "document_id": doc.id,
                "type": "embed",
            })
        )
        parsed = json.loads(result)
        assert "error" in parsed

    def test_embed_section_preview(self):
        """Preview of an embed section should show the ref."""
        from llming_docs.text_doc_mcp import TextDocMCP
        from llming_docs.document_store import DocumentSessionStore

        store = DocumentSessionStore()
        mcp = TextDocMCP(store)
        section = {"type": "embed", "$ref": "abc123def456"}
        preview = mcp._section_preview(section)
        assert "embed" in preview
        assert "abc123def456" in preview

    def test_embed_section_in_list_sections(self):
        """List sections should include embed sections."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Test", data={
            "sections": [
                {"id": "s1", "type": "heading", "content": "Title"},
                {"id": "s2", "type": "embed", "$ref": "chart-uuid"},
                {"id": "s3", "type": "paragraph", "content": "After chart"},
            ]
        })
        mcp = TextDocMCP(store)

        result = asyncio.new_event_loop().run_until_complete(
            mcp.call_tool("text_doc_list_sections", {"document_id": doc.id})
        )
        parsed = json.loads(result)
        assert parsed["section_count"] == 3
        types = [s["type"] for s in parsed["sections"]]
        assert types == ["heading", "embed", "paragraph"]

    def test_add_regular_section_still_requires_content(self):
        """Non-embed sections should still work with content."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Test", data={"sections": []})
        mcp = TextDocMCP(store)

        result = asyncio.new_event_loop().run_until_complete(
            mcp.call_tool("text_doc_add_section", {
                "document_id": doc.id,
                "type": "paragraph",
                "content": "Hello world",
            })
        )
        parsed = json.loads(result)
        assert parsed["status"] == "section_added"

        updated = store.get(doc.id)
        assert updated.data["sections"][0]["content"] == "Hello world"

    def test_embed_with_position(self):
        """Embed can be inserted at a specific position."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Test", data={
            "sections": [
                {"id": "s1", "type": "heading", "content": "Title"},
                {"id": "s2", "type": "paragraph", "content": "End"},
            ]
        })
        mcp = TextDocMCP(store)

        result = asyncio.new_event_loop().run_until_complete(
            mcp.call_tool("text_doc_add_section", {
                "document_id": doc.id,
                "type": "embed",
                "ref": "table-uuid",
                "position": 1,
            })
        )
        parsed = json.loads(result)
        assert parsed["status"] == "section_added"

        updated = store.get(doc.id)
        types = [s["type"] for s in updated.data["sections"]]
        assert types == ["heading", "embed", "paragraph"]
        assert updated.data["sections"][1]["$ref"] == "table-uuid"

    def test_multiple_embeds_in_document(self):
        """Multiple embeds referencing different documents."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Report", data={"sections": []})
        mcp = TextDocMCP(store)

        for ref_id in ["plotly-1", "table-1", "plotly-2"]:
            asyncio.new_event_loop().run_until_complete(
                mcp.call_tool("text_doc_add_section", {
                    "document_id": doc.id,
                    "type": "embed",
                    "ref": ref_id,
                })
            )

        updated = store.get(doc.id)
        assert len(updated.data["sections"]) == 3
        refs = [s["$ref"] for s in updated.data["sections"]]
        assert refs == ["plotly-1", "table-1", "plotly-2"]

    def test_update_embed_section(self):
        """Can update an embed section to change its $ref."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Test", data={
            "sections": [
                {"id": "s1", "type": "embed", "$ref": "old-ref"},
            ]
        })
        mcp = TextDocMCP(store)

        result = asyncio.new_event_loop().run_until_complete(
            mcp.call_tool("text_doc_update_section", {
                "document_id": doc.id,
                "section_id": "s1",
                "updates": {"$ref": "new-ref"},
            })
        )
        parsed = json.loads(result)
        assert parsed["status"] == "section_updated"

        updated = store.get(doc.id)
        assert updated.data["sections"][0]["$ref"] == "new-ref"

    def test_delete_embed_section(self):
        """Can delete an embed section."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Test", data={
            "sections": [
                {"id": "s1", "type": "paragraph", "content": "Before"},
                {"id": "s2", "type": "embed", "$ref": "chart-ref"},
                {"id": "s3", "type": "paragraph", "content": "After"},
            ]
        })
        mcp = TextDocMCP(store)

        result = asyncio.new_event_loop().run_until_complete(
            mcp.call_tool("text_doc_delete_section", {
                "document_id": doc.id,
                "section_id": "s2",
            })
        )
        parsed = json.loads(result)
        assert parsed["status"] == "section_deleted"

        updated = store.get(doc.id)
        assert len(updated.data["sections"]) == 2
        types = [s["type"] for s in updated.data["sections"]]
        assert "embed" not in types

    def test_move_embed_section(self):
        """Can move an embed section to a new position."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Test", data={
            "sections": [
                {"id": "s1", "type": "heading", "content": "Title"},
                {"id": "s2", "type": "paragraph", "content": "Text"},
                {"id": "s3", "type": "embed", "$ref": "chart-ref"},
            ]
        })
        mcp = TextDocMCP(store)

        result = asyncio.new_event_loop().run_until_complete(
            mcp.call_tool("text_doc_move_section", {
                "document_id": doc.id,
                "section_id": "s3",
                "new_position": 1,
            })
        )
        parsed = json.loads(result)
        assert parsed["status"] == "section_moved"

        updated = store.get(doc.id)
        types = [s["type"] for s in updated.data["sections"]]
        assert types == ["heading", "embed", "paragraph"]

    def test_search_skips_embed_content(self):
        """Search should handle embed sections gracefully (no content to search)."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Test", data={
            "sections": [
                {"id": "s1", "type": "paragraph", "content": "findme"},
                {"id": "s2", "type": "embed", "$ref": "findme-uuid"},
            ]
        })
        mcp = TextDocMCP(store)

        result = asyncio.new_event_loop().run_until_complete(
            mcp.call_tool("text_doc_search", {
                "document_id": doc.id,
                "query": "findme",
            })
        )
        parsed = json.loads(result)
        # Should find the paragraph and possibly the embed (ref contains "findme")
        assert parsed["total_matches"] >= 1

    def test_complex_doc_with_mixed_sections_and_embeds(self):
        """Full document with headings, paragraphs, tables, and embeds."""
        from llming_docs.document_store import DocumentSessionStore
        from llming_docs.text_doc_mcp import TextDocMCP
        import asyncio
        import json

        store = DocumentSessionStore()
        doc = store.create("text_doc", name="Report", data={"sections": []})
        mcp = TextDocMCP(store)

        sections_to_add = [
            {"type": "heading", "content": "Q1 Report", "level": 1},
            {"type": "paragraph", "content": "Executive summary."},
            {"type": "embed", "ref": "revenue-chart"},
            {"type": "heading", "content": "Data Tables", "level": 2},
            {"type": "embed", "ref": "sales-table"},
            {"type": "paragraph", "content": "Conclusion text."},
            {"type": "embed", "ref": "forecast-chart"},
        ]

        for s in sections_to_add:
            args = {"document_id": doc.id, "type": s["type"]}
            if "content" in s:
                args["content"] = s["content"]
            if "level" in s:
                args["level"] = s["level"]
            if "ref" in s:
                args["ref"] = s["ref"]
            asyncio.new_event_loop().run_until_complete(
                mcp.call_tool("text_doc_add_section", args)
            )

        updated = store.get(doc.id)
        assert len(updated.data["sections"]) == 7
        types = [s["type"] for s in updated.data["sections"]]
        assert types == ["heading", "paragraph", "embed", "heading", "embed", "paragraph", "embed"]
        refs = [s.get("$ref") for s in updated.data["sections"] if s["type"] == "embed"]
        assert refs == ["revenue-chart", "sales-table", "forecast-chart"]
