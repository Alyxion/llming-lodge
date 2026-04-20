"""End-to-end integration test for the embed pipeline.

Simulates the full flow:
1. Create a plotly chart via MathMCP
2. Create Word/PPTX/Table documents that embed the chart via $ref
3. Export each to binary format (DOCX/PPTX/XLSX) with the chart rendered as PNG
4. Create an email draft attaching all three documents
5. Verify the email attachments contain valid binary files

No real LLM or MS Graph — all MCP tools are called directly,
and the MS Graph is mocked.
"""

import asyncio
import base64
import io
import json
import zipfile

import pytest

from llming_docs.document_store import DocumentSessionStore
from llming_docs.text_doc_mcp import TextDocMCP
from llming_docs.email_mcp import EmailDraftMCP
from llming_docs.table_exporter import export_xlsx
from llming_docs.html_exporter import export_html
from llming_docs.word_exporter import export_docx
from llming_docs.render import (
    EMBED_BEHAVIOR,
    EmbedBehavior,
    RenderContext,
    can_embed,
    get_embed_behavior,
    render_to,
)


# ── Helpers ───────────────────────────────────────────────────────────────

# Valid 4×4 red RGBA PNG (python-docx requires a parseable PNG with IHDR+IDAT+IEND)
_RED_PNG_B64 = (
    "iVBORw0KGgoAAAANSUhEUgAAAAQAAAAECAYAAACp8Z5+AAAAEklEQVR4nGP4"
    "z8DwHxkzkC4AADxAH+HggXe0AAAAAElFTkSuQmCC"
)
_RED_PNG_DATA_URI = f"data:image/png;base64,{_RED_PNG_B64}"
_RED_PNG_BYTES = base64.b64decode(_RED_PNG_B64)


def _run(coro):
    """Run an async coroutine in a fresh event loop."""
    loop = asyncio.new_event_loop()
    try:
        return loop.run_until_complete(coro)
    finally:
        loop.close()


def _is_valid_zip(data: bytes) -> bool:
    """Check if bytes are a valid ZIP archive (DOCX/PPTX/XLSX are all ZIP)."""
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            return len(zf.namelist()) > 0
    except zipfile.BadZipFile:
        return False


def _is_valid_docx(data: bytes) -> bool:
    """Check if bytes are a valid DOCX file."""
    if not _is_valid_zip(data):
        return False
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return "word/document.xml" in zf.namelist()


def _is_valid_xlsx(data: bytes) -> bool:
    """Check if bytes are a valid XLSX file."""
    if not _is_valid_zip(data):
        return False
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        names = zf.namelist()
        return "xl/workbook.xml" in names or "xl/worksheets/sheet1.xml" in names


def _docx_has_image(data: bytes) -> bool:
    """Check if DOCX contains at least one image."""
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        return any("media/image" in n for n in zf.namelist())


def _create_plotly_chart(store: DocumentSessionStore) -> str:
    """Create a plotly chart document (simulates math_plot_2d output).

    Returns the document ID.
    """
    import math
    x = [i * 0.1 for i in range(-100, 101)]
    y = [math.sin(v) for v in x]

    plotly_data = {
        "data": [
            {
                "x": x,
                "y": y,
                "type": "scatter",
                "mode": "lines",
                "name": "sin(x)",
                "line": {"color": "#6366f1", "width": 2},
            }
        ],
        "layout": {
            "title": {"text": "sin(x)"},
            "xaxis": {"title": {"text": "x"}},
            "yaxis": {"title": {"text": "y"}},
            "showlegend": True,
        },
    }
    doc = store.create(type="plotly", name="sin(x) Plot", data=plotly_data, skip_validation=True)
    return doc.id


# ═══════════════════════════════════════════════════════════════════════════
# STEP 1: Create chart and embed it in a Word document
# ═══════════════════════════════════════════════════════════════════════════


class TestStep1_EmbedInWordDoc:
    """Create a plotly chart, then create a Word doc that embeds it via $ref."""

    def test_create_chart_and_embed_in_text_doc(self):
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)

        # Verify chart exists in store
        chart_doc = store.get(chart_id)
        assert chart_doc is not None
        assert chart_doc.type == "plotly"
        assert len(chart_doc.data["data"]) == 1

        # Create Word document with embed section referencing chart
        mcp = TextDocMCP(store)
        word_doc = store.create(type="text_doc", name="Analysis Report", data={"sections": []}, skip_validation=True)

        # Add heading
        result = _run(mcp.call_tool("text_doc_add_section", {
            "document_id": word_doc.id,
            "type": "heading",
            "content": "Sine Function Analysis",
            "level": 1,
        }))
        assert json.loads(result)["status"] == "section_added"

        # Add paragraph
        result = _run(mcp.call_tool("text_doc_add_section", {
            "document_id": word_doc.id,
            "type": "paragraph",
            "content": "Below is the plot of sin(x):",
        }))
        assert json.loads(result)["status"] == "section_added"

        # Add embed section referencing the chart
        result = _run(mcp.call_tool("text_doc_add_section", {
            "document_id": word_doc.id,
            "type": "embed",
            "ref": chart_id,
        }))
        parsed = json.loads(result)
        assert parsed["status"] == "section_added"
        assert parsed["section_count"] == 3

        # Verify the embed section has $ref
        updated = store.get(word_doc.id)
        sections = updated.data["sections"]
        assert sections[2]["type"] == "embed"
        assert sections[2]["$ref"] == chart_id

    def test_embed_behavior_lookup_for_plotly(self):
        """Verify plotly is registered as graphic mode for embedding."""
        behavior = get_embed_behavior("plotly")
        assert behavior is not None
        assert behavior.mode == "graphic"
        assert behavior.aspect == 1.6


# ═══════════════════════════════════════════════════════════════════════════
# STEP 2: Export Word doc to DOCX with embedded chart rendered as PNG
# ═══════════════════════════════════════════════════════════════════════════


class TestStep2_ExportDocxWithEmbeddedChart:
    """Export a Word doc with an embed section → DOCX with chart as image."""

    def test_export_docx_with_chart_image(self):
        """Simulate client-side embed resolution + server-side DOCX export.

        The client resolves embed→image before sending to the server.
        This test simulates that flow.
        """
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)

        # Build a Word doc spec as it would look AFTER client-side resolution:
        # The embed section has been replaced with an image section
        # (client rendered plotly→PNG before sending to server)
        export_spec = {
            "title": "Analysis Report",
            "sections": [
                {"type": "heading", "level": 1, "content": "Sine Function Analysis"},
                {"type": "paragraph", "content": "Below is the plot of sin(x):"},
                # Client resolved embed→image:
                {"type": "image", "data": _RED_PNG_DATA_URI},
                {"type": "paragraph", "content": "The function oscillates between -1 and 1."},
            ],
        }

        docx_bytes = export_docx(export_spec)
        assert _is_valid_docx(docx_bytes)
        assert _docx_has_image(docx_bytes)

    def test_export_docx_with_chart_images_dict(self):
        """Export using chart_images dict (alternative flow)."""
        export_spec = {
            "title": "Chart Report",
            "sections": [
                {"type": "heading", "level": 1, "content": "Chart Report"},
                {"type": "chart", "_chartImageId": "chart_0"},
            ],
        }
        chart_images = {"chart_0": _RED_PNG_DATA_URI}
        docx_bytes = export_docx(export_spec, chart_images=chart_images)
        assert _is_valid_docx(docx_bytes)
        assert _docx_has_image(docx_bytes)

    def test_render_to_docx(self):
        """Use render_to() API for text_doc → DOCX."""
        spec = {
            "title": "Report",
            "sections": [
                {"type": "paragraph", "content": "Text content"},
                {"type": "image", "data": _RED_PNG_DATA_URI},
            ],
        }
        result = render_to("text_doc", spec, "docx")
        assert result.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert _is_valid_docx(result.data)
        assert _docx_has_image(result.data)


# ═══════════════════════════════════════════════════════════════════════════
# STEP 3: Create table document and export to XLSX
# ═══════════════════════════════════════════════════════════════════════════


class TestStep3_TableToXlsx:
    """Create a table document and export to XLSX."""

    def test_table_export_via_render_to(self):
        store = DocumentSessionStore()
        table_data = {
            "columns": ["x", "sin(x)", "cos(x)"],
            "rows": [
                [0, 0.0, 1.0],
                [1.57, 1.0, 0.0],
                [3.14, 0.0, -1.0],
                [4.71, -1.0, 0.0],
                [6.28, 0.0, 1.0],
            ],
        }
        table_doc = store.create(type="table", name="Trig Values", data=table_data, skip_validation=True)

        result = render_to("table", table_data, "xlsx")
        assert result.content_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        assert _is_valid_xlsx(result.data)

        # Verify content
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(result.data), read_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(values_only=True))
        assert rows[0] == ("x", "sin(x)", "cos(x)")
        assert rows[1][0] == 0  # first data row

    def test_table_with_chart_embed_in_word(self):
        """Word doc with both a chart embed and a table section."""
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)
        mcp = TextDocMCP(store)

        word_doc = store.create(type="text_doc", name="Full Report", data={"sections": []}, skip_validation=True)

        # Add embed for chart
        _run(mcp.call_tool("text_doc_add_section", {
            "document_id": word_doc.id,
            "type": "embed",
            "ref": chart_id,
        }))

        # Add a native table section
        _run(mcp.call_tool("text_doc_add_section", {
            "document_id": word_doc.id,
            "type": "table",
            "content": [["x", "sin(x)"], ["0", "0"], ["π/2", "1"]],
        }))

        updated = store.get(word_doc.id)
        sections = updated.data["sections"]
        assert len(sections) == 2
        assert sections[0]["type"] == "embed"
        assert sections[0]["$ref"] == chart_id
        assert sections[1]["type"] == "table"


# ═══════════════════════════════════════════════════════════════════════════
# STEP 4: Simulate client-side embed resolution for all export formats
# ═══════════════════════════════════════════════════════════════════════════


class TestStep4_ClientSideEmbedResolution:
    """Simulate what the client does: resolve embeds, then export.

    The client looks up the referenced document, checks its type against
    EMBED_BEHAVIOR, and converts:
    - graphic → image (PNG data URI)
    - table → native table section
    - text → inline sections
    """

    def _resolve_embed_for_docx(self, store, embed_section):
        """Simulate client-side embed resolution for DOCX export.

        Returns a list of resolved sections to replace the embed.
        """
        ref_id = embed_section.get("$ref")
        if not ref_id:
            return [embed_section]

        doc = store.get(ref_id)
        if not doc:
            return [{"type": "paragraph", "content": f"[Missing document: {ref_id}]"}]

        behavior = get_embed_behavior(doc.type)
        if not behavior:
            return [{"type": "paragraph", "content": f"[Unknown type: {doc.type}]"}]

        if behavior.mode == "graphic":
            # Client would render to PNG via Plotly.toImage()
            # We simulate with our test PNG
            return [{"type": "image", "data": _RED_PNG_DATA_URI}]

        elif behavior.mode == "table":
            data = doc.data or {}
            headers = data.get("headers") or data.get("columns") or []
            col_labels = [
                (h.get("label", h.get("key", "")) if isinstance(h, dict) else str(h))
                for h in headers
            ]
            col_keys = [
                (h.get("key", h.get("label", "")) if isinstance(h, dict) else str(h))
                for h in headers
            ]
            rows = data.get("rows", [])
            flat_rows = []
            for row in rows:
                if isinstance(row, dict):
                    flat_rows.append([row.get(k, "") for k in col_keys])
                elif isinstance(row, (list, tuple)):
                    flat_rows.append(list(row))
                else:
                    flat_rows.append([row])
            return [{"type": "table", "headers": col_labels, "rows": flat_rows}]

        elif behavior.mode == "text" and isinstance(doc.data, dict):
            return doc.data.get("sections", [])

        return [{"type": "paragraph", "content": f"[Unresolvable: {doc.type}]"}]

    def test_resolve_plotly_embed_to_image(self):
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)

        resolved = self._resolve_embed_for_docx(store, {"$ref": chart_id})
        assert len(resolved) == 1
        assert resolved[0]["type"] == "image"
        assert resolved[0]["data"].startswith("data:image/png")

    def test_resolve_table_embed_to_native_table(self):
        store = DocumentSessionStore()
        table_doc = store.create(type="table", name="Data", data={
            "columns": ["A", "B"],
            "rows": [["1", "2"], ["3", "4"]],
        }, skip_validation=True)

        resolved = self._resolve_embed_for_docx(store, {"$ref": table_doc.id})
        assert len(resolved) == 1
        assert resolved[0]["type"] == "table"
        assert resolved[0]["headers"] == ["A", "B"]
        assert resolved[0]["rows"] == [["1", "2"], ["3", "4"]]

    def test_resolve_text_doc_embed_to_sections(self):
        store = DocumentSessionStore()
        sub_doc = store.create(type="text_doc", name="Sub", data={
            "sections": [
                {"type": "paragraph", "content": "Inline text from sub-doc"},
            ],
        }, skip_validation=True)

        resolved = self._resolve_embed_for_docx(store, {"$ref": sub_doc.id})
        assert len(resolved) == 1
        assert resolved[0]["type"] == "paragraph"
        assert "sub-doc" in resolved[0]["content"]

    def test_full_docx_with_resolved_embeds(self):
        """Build a complete Word doc spec with resolved embeds → export to DOCX."""
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)
        table_doc = store.create(type="table", name="Values", data={
            "columns": ["x", "y"],
            "rows": [["0", "0"], ["1", "0.84"]],
        }, skip_validation=True)

        # Simulate what client does: resolve each embed section
        raw_sections = [
            {"type": "heading", "level": 1, "content": "Report"},
            {"type": "paragraph", "content": "Chart below:"},
            {"type": "embed", "$ref": chart_id},
            {"type": "paragraph", "content": "Table below:"},
            {"type": "embed", "$ref": table_doc.id},
        ]

        resolved_sections = []
        for s in raw_sections:
            if s["type"] == "embed":
                resolved_sections.extend(self._resolve_embed_for_docx(store, s))
            else:
                resolved_sections.append(s)

        export_spec = {"title": "Report", "sections": resolved_sections}
        docx_bytes = export_docx(export_spec)
        assert _is_valid_docx(docx_bytes)
        assert _docx_has_image(docx_bytes)

        # Verify the table made it
        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(docx_bytes))
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "x"


# ═══════════════════════════════════════════════════════════════════════════
# STEP 5: Email draft with all three document types as attachments
# ═══════════════════════════════════════════════════════════════════════════


class TestStep5_EmailDraftWithAttachments:
    """Create email draft referencing Word, Table, and Plotly documents."""

    def test_create_email_draft_with_attachments(self):
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)

        # Create Word doc with embed
        word_doc = store.create(type="text_doc", name="Report.docx", data={
            "sections": [
                {"type": "heading", "level": 1, "content": "Report"},
                {"type": "embed", "$ref": chart_id},
            ],
        }, skip_validation=True)

        # Create table doc
        table_doc = store.create(type="table", name="Data.xlsx", data={
            "columns": ["x", "sin(x)"],
            "rows": [[0, 0.0], [1.57, 1.0], [3.14, 0.0]],
        }, skip_validation=True)

        # Create email draft
        email_mcp = EmailDraftMCP(store)
        email_doc = store.create(type="email_draft", name="Report Email", data={
            "subject": "Monthly Report",
            "to": ["test@example.com"],
            "cc": [],
            "bcc": [],
            "body_html": "<p>Please find the attached report.</p>",
            "attachments": [],
        }, skip_validation=True)

        # Add Word doc as attachment
        result = _run(email_mcp.call_tool("email_add_attachment", {
            "document_id": email_doc.id,
            "ref": word_doc.id,
            "name": "Report.docx",
        }))
        assert json.loads(result)["status"] == "attachment_added"

        # Add table as attachment
        result = _run(email_mcp.call_tool("email_add_attachment", {
            "document_id": email_doc.id,
            "ref": table_doc.id,
            "name": "Data.xlsx",
        }))
        assert json.loads(result)["status"] == "attachment_added"

        # Add chart directly as attachment
        result = _run(email_mcp.call_tool("email_add_attachment", {
            "document_id": email_doc.id,
            "ref": chart_id,
            "name": "Chart.png",
        }))
        assert json.loads(result)["status"] == "attachment_added"

        # Verify draft has 3 attachments
        updated = store.get(email_doc.id)
        attachments = updated.data["attachments"]
        assert len(attachments) == 3
        refs = [a["ref"] for a in attachments]
        assert word_doc.id in refs
        assert table_doc.id in refs
        assert chart_id in refs


# ═══════════════════════════════════════════════════════════════════════════
# STEP 6: Simulate attachment resolution and export (client-side flow)
# ═══════════════════════════════════════════════════════════════════════════


class TestStep6_ResolveAndExportAttachments:
    """Simulate the client's _resolveAttachments flow.

    For each attachment ref:
    - Look up document in store
    - Export to binary format based on type
    - Verify the binary is valid
    """

    def _resolve_attachment(self, store, ref_id):
        """Simulate client _attachmentExporters[entry.lang](entry, name).

        Returns (name, content_type, data_bytes) or None.
        """
        doc = store.get(ref_id)
        if not doc:
            return None

        if doc.type == "plotly":
            # Client renders to PNG via Plotly.toImage()
            # We simulate with test PNG
            return ("chart.png", "image/png", _RED_PNG_BYTES)

        elif doc.type == "table":
            xlsx_bytes = export_xlsx(doc.data)
            return ("data.xlsx",
                    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    xlsx_bytes)

        elif doc.type == "text_doc":
            # Resolve embeds in sections before exporting
            resolved_sections = []
            for s in (doc.data or {}).get("sections", []):
                if s.get("type") == "embed" and s.get("$ref"):
                    ref_doc = store.get(s["$ref"])
                    if ref_doc and get_embed_behavior(ref_doc.type):
                        b = get_embed_behavior(ref_doc.type)
                        if b.mode == "graphic":
                            resolved_sections.append({"type": "image", "data": _RED_PNG_DATA_URI})
                        elif b.mode == "table":
                            data = ref_doc.data or {}
                            resolved_sections.append({
                                "type": "table",
                                "headers": data.get("columns", []),
                                "rows": data.get("rows", []),
                            })
                        else:
                            resolved_sections.extend((ref_doc.data or {}).get("sections", []))
                    continue
                resolved_sections.append(s)

            export_spec = dict(doc.data)
            export_spec["sections"] = resolved_sections
            docx_bytes = export_docx(export_spec)
            return ("report.docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    docx_bytes)

        elif doc.type == "html" or doc.type == "html_sandbox":
            html_bytes = export_html(doc.data)
            return ("page.html", "text/html", html_bytes)

        return None

    def test_resolve_all_attachment_types(self):
        """Full pipeline: chart + table + word doc → resolved binary attachments."""
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)

        # Word doc with embedded chart
        word_doc = store.create(type="text_doc", name="Report", data={
            "title": "Analysis",
            "sections": [
                {"type": "heading", "level": 1, "content": "Analysis"},
                {"type": "embed", "$ref": chart_id},
                {"type": "paragraph", "content": "Conclusion."},
            ],
        }, skip_validation=True)

        # Table doc
        table_doc = store.create(type="table", name="Data", data={
            "columns": ["x", "sin(x)"],
            "rows": [[0, 0], [1, 0.84], [2, 0.91]],
        }, skip_validation=True)

        # Resolve each attachment
        chart_att = self._resolve_attachment(store, chart_id)
        assert chart_att is not None
        name, ctype, data = chart_att
        assert ctype == "image/png"
        assert len(data) > 0

        table_att = self._resolve_attachment(store, table_doc.id)
        assert table_att is not None
        name, ctype, data = table_att
        assert ctype == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        assert _is_valid_xlsx(data)

        word_att = self._resolve_attachment(store, word_doc.id)
        assert word_att is not None
        name, ctype, data = word_att
        assert ctype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert _is_valid_docx(data)
        # The DOCX should contain the chart as an embedded image
        assert _docx_has_image(data)

    def test_word_doc_with_table_embed_exports_native_table(self):
        """Word doc with embedded table → DOCX has a native Word table."""
        store = DocumentSessionStore()
        table_doc = store.create(type="table", name="Metrics", data={
            "columns": ["Metric", "Value"],
            "rows": [["Revenue", "$1.2M"], ["Growth", "+15%"]],
        }, skip_validation=True)

        word_doc = store.create(type="text_doc", name="Report", data={
            "title": "Metrics Report",
            "sections": [
                {"type": "heading", "level": 1, "content": "Key Metrics"},
                {"type": "embed", "$ref": table_doc.id},
            ],
        }, skip_validation=True)

        att = self._resolve_attachment(store, word_doc.id)
        assert att is not None
        _, _, docx_bytes = att
        assert _is_valid_docx(docx_bytes)

        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(docx_bytes))
        assert len(doc.tables) == 1
        assert doc.tables[0].cell(0, 0).text == "Metric"
        assert doc.tables[0].cell(1, 0).text == "Revenue"


# ═══════════════════════════════════════════════════════════════════════════
# STEP 7: Full email send simulation with mocked MS Graph
# ═══════════════════════════════════════════════════════════════════════════


class TestStep7_EmailSendWithMockedGraph:
    """Simulate the full email send flow with mocked MS Graph.

    Creates chart → word doc (with embed) → table → email draft → resolves
    attachments → simulates send via mocked Graph API → verifies payload.
    """

    def test_full_email_pipeline(self):
        """End-to-end: chart → docs → email draft → resolve → mock send."""
        store = DocumentSessionStore()

        # 1. Create chart
        chart_id = _create_plotly_chart(store)

        # 2. Create Word doc with embedded chart
        mcp = TextDocMCP(store)
        word_doc = store.create(type="text_doc", name="Analysis", data={
            "title": "Sin Analysis",
            "sections": [],
        }, skip_validation=True)
        _run(mcp.call_tool("text_doc_add_section", {
            "document_id": word_doc.id,
            "type": "heading",
            "content": "Sine Analysis",
            "level": 1,
        }))
        _run(mcp.call_tool("text_doc_add_section", {
            "document_id": word_doc.id,
            "type": "embed",
            "ref": chart_id,
        }))
        _run(mcp.call_tool("text_doc_add_section", {
            "document_id": word_doc.id,
            "type": "paragraph",
            "content": "The sine function oscillates.",
        }))

        # 3. Create table doc
        table_doc = store.create(type="table", name="Values", data={
            "columns": ["x", "sin(x)"],
            "rows": [[0, 0], [1.57, 1.0], [3.14, 0], [4.71, -1.0]],
        }, skip_validation=True)

        # 4. Create email draft with all attachments
        email_mcp = EmailDraftMCP(store)
        email_doc = store.create(type="email_draft", name="Report Email", data={
            "subject": "Monthly Sine Report",
            "to": ["recipient@example.com"],
            "cc": ["cc@example.com"],
            "bcc": [],
            "body_html": "<h1>Report</h1><p>See attachments.</p>",
            "attachments": [],
        }, skip_validation=True)

        for ref_id, name in [
            (word_doc.id, "Analysis.docx"),
            (table_doc.id, "Values.xlsx"),
            (chart_id, "Chart.png"),
        ]:
            result = _run(email_mcp.call_tool("email_add_attachment", {
                "document_id": email_doc.id,
                "ref": ref_id,
                "name": name,
            }))
            assert json.loads(result)["status"] == "attachment_added"

        # 5. Resolve attachments (simulate client _resolveAttachments)
        email_data = store.get(email_doc.id).data
        resolved_attachments = []
        for att in email_data["attachments"]:
            ref_doc = store.get(att["ref"])
            assert ref_doc is not None, f"Missing doc for ref {att['ref']}"

            # Simulate export based on type
            if ref_doc.type == "plotly":
                resolved_attachments.append({
                    "type": "file",
                    "name": att["name"],
                    "content_type": "image/png",
                    "data": _RED_PNG_B64,
                    "size": len(_RED_PNG_BYTES),
                })
            elif ref_doc.type == "table":
                xlsx = export_xlsx(ref_doc.data)
                resolved_attachments.append({
                    "type": "file",
                    "name": att["name"],
                    "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    "data": base64.b64encode(xlsx).decode(),
                    "size": len(xlsx),
                })
            elif ref_doc.type == "text_doc":
                # Resolve embeds before exporting
                sections = []
                for s in (ref_doc.data or {}).get("sections", []):
                    if s.get("type") == "embed":
                        sections.append({"type": "image", "data": _RED_PNG_DATA_URI})
                    else:
                        sections.append(s)
                spec = dict(ref_doc.data)
                spec["sections"] = sections
                docx = export_docx(spec)
                resolved_attachments.append({
                    "type": "file",
                    "name": att["name"],
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "data": base64.b64encode(docx).decode(),
                    "size": len(docx),
                })

        # 6. Simulate MS Graph send payload
        assert len(resolved_attachments) == 3

        # Build Graph API payload (simulates what OfficeMailHandler does)
        graph_attachments = []
        for att in resolved_attachments:
            graph_attachments.append({
                "@odata.type": "#microsoft.graph.fileAttachment",
                "name": att["name"],
                "contentType": att["content_type"],
                "contentBytes": att["data"],
            })

        graph_payload = {
            "message": {
                "subject": email_data["subject"],
                "body": {"contentType": "HTML", "content": email_data["body_html"]},
                "toRecipients": [{"emailAddress": {"address": a}} for a in email_data["to"]],
                "ccRecipients": [{"emailAddress": {"address": a}} for a in email_data.get("cc", [])],
            },
            "saveToSentItems": "true",
            "attachments": graph_attachments,
        }

        # 7. Verify the payload
        assert graph_payload["message"]["subject"] == "Monthly Sine Report"
        assert len(graph_payload["attachments"]) == 3

        # Verify each attachment is valid binary
        for att in graph_payload["attachments"]:
            raw = base64.b64decode(att["contentBytes"])
            assert len(raw) > 0

            if att["name"].endswith(".docx"):
                assert _is_valid_docx(raw), f"{att['name']} is not valid DOCX"
                assert _docx_has_image(raw), f"{att['name']} has no embedded image"

            elif att["name"].endswith(".xlsx"):
                assert _is_valid_xlsx(raw), f"{att['name']} is not valid XLSX"

            elif att["name"].endswith(".png"):
                # PNG magic bytes
                assert raw[:4] == b'\x89PNG', f"{att['name']} is not valid PNG"

    def test_email_draft_get_returns_attachments(self):
        """Verify email_get_draft returns attachment refs."""
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)

        email_mcp = EmailDraftMCP(store)
        email_doc = store.create(type="email_draft", name="Test", data={
            "subject": "Test",
            "to": ["a@b.com"],
            "body_html": "<p>Hi</p>",
            "attachments": [],
        }, skip_validation=True)

        _run(email_mcp.call_tool("email_add_attachment", {
            "document_id": email_doc.id,
            "ref": chart_id,
            "name": "Plot.png",
        }))

        result = _run(email_mcp.call_tool("email_get_draft", {
            "document_id": email_doc.id,
        }))
        parsed = json.loads(result)
        assert parsed["subject"] == "Test"
        assert len(parsed["attachments"]) == 1
        assert parsed["attachments"][0]["ref"] == chart_id

    def test_duplicate_attachment_rejected(self):
        """Adding same ref twice should fail."""
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)

        email_mcp = EmailDraftMCP(store)
        email_doc = store.create(type="email_draft", name="Test", data={
            "subject": "Test",
            "to": ["a@b.com"],
            "body_html": "",
            "attachments": [],
        }, skip_validation=True)

        _run(email_mcp.call_tool("email_add_attachment", {
            "document_id": email_doc.id,
            "ref": chart_id,
            "name": "Chart",
        }))

        result = _run(email_mcp.call_tool("email_add_attachment", {
            "document_id": email_doc.id,
            "ref": chart_id,
            "name": "Chart Again",
        }))
        parsed = json.loads(result)
        assert parsed["status"] == "already_attached"


# ═══════════════════════════════════════════════════════════════════════════
# STEP 8: Edge cases and deeply nested scenarios
# ═══════════════════════════════════════════════════════════════════════════


class TestStep8_EdgeCases:
    """Wild edge cases for the embed pipeline."""

    def test_word_doc_with_multiple_embeds_of_same_chart(self):
        """Same chart embedded twice in one Word doc."""
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)

        mcp = TextDocMCP(store)
        doc = store.create(type="text_doc", name="Double Chart", data={"sections": []}, skip_validation=True)

        _run(mcp.call_tool("text_doc_add_section", {
            "document_id": doc.id, "type": "embed", "ref": chart_id,
        }))
        _run(mcp.call_tool("text_doc_add_section", {
            "document_id": doc.id, "type": "paragraph", "content": "Between charts",
        }))
        _run(mcp.call_tool("text_doc_add_section", {
            "document_id": doc.id, "type": "embed", "ref": chart_id,
        }))

        updated = store.get(doc.id)
        assert len(updated.data["sections"]) == 3
        assert updated.data["sections"][0]["$ref"] == chart_id
        assert updated.data["sections"][2]["$ref"] == chart_id

        # Export with resolved embeds
        sections = []
        for s in updated.data["sections"]:
            if s["type"] == "embed":
                sections.append({"type": "image", "data": _RED_PNG_DATA_URI})
            else:
                sections.append(s)

        docx = export_docx({"title": "Double", "sections": sections})
        assert _is_valid_docx(docx)
        assert _docx_has_image(docx)

    def test_embed_missing_ref_handled_gracefully(self):
        """Embed with nonexistent ref should not crash export."""
        store = DocumentSessionStore()
        mcp = TextDocMCP(store)
        doc = store.create(type="text_doc", name="Test", data={"sections": []}, skip_validation=True)

        _run(mcp.call_tool("text_doc_add_section", {
            "document_id": doc.id,
            "type": "embed",
            "ref": "nonexistent-uuid",
        }))

        updated = store.get(doc.id)
        assert updated.data["sections"][0]["$ref"] == "nonexistent-uuid"

        # On export, the embed with missing ref would be skipped/placeholder
        sections = []
        for s in updated.data["sections"]:
            if s["type"] == "embed":
                ref_doc = store.get(s.get("$ref", ""))
                if ref_doc:
                    sections.append({"type": "image", "data": _RED_PNG_DATA_URI})
                else:
                    sections.append({"type": "paragraph", "content": "[Missing reference]"})
            else:
                sections.append(s)

        docx = export_docx({"title": "Test", "sections": sections})
        assert _is_valid_docx(docx)

    def test_embed_chain_text_doc_in_text_doc(self):
        """Text doc embedding another text doc (text mode)."""
        store = DocumentSessionStore()

        inner_doc = store.create(type="text_doc", name="Inner", data={
            "sections": [
                {"type": "paragraph", "content": "Inner paragraph 1"},
                {"type": "paragraph", "content": "Inner paragraph 2"},
            ],
        }, skip_validation=True)

        outer_doc = store.create(type="text_doc", name="Outer", data={
            "sections": [
                {"type": "heading", "level": 1, "content": "Outer Doc"},
                {"type": "embed", "$ref": inner_doc.id},
                {"type": "paragraph", "content": "After inner doc"},
            ],
        }, skip_validation=True)

        # Resolve: text embed → splice inner sections
        behavior = get_embed_behavior("text_doc")
        assert behavior.mode == "text"

        sections = []
        for s in outer_doc.data["sections"]:
            if s["type"] == "embed":
                ref = store.get(s["$ref"])
                if ref and get_embed_behavior(ref.type).mode == "text":
                    sections.extend(ref.data.get("sections", []))
                    continue
            sections.append(s)

        docx = export_docx({"title": "Nested", "sections": sections})
        assert _is_valid_docx(docx)

        from docx import Document as DocxDoc
        doc = DocxDoc(io.BytesIO(docx))
        texts = [p.text for p in doc.paragraphs]
        assert "Inner paragraph 1" in texts
        assert "Inner paragraph 2" in texts
        assert "After inner doc" in texts

    def test_email_with_all_document_types(self):
        """Email with plotly + table + word + html attachments."""
        store = DocumentSessionStore()
        chart_id = _create_plotly_chart(store)

        table_doc = store.create(type="table", name="T", data={
            "columns": ["A"], "rows": [["1"]],
        }, skip_validation=True)
        word_doc = store.create(type="text_doc", name="W", data={
            "title": "W", "sections": [{"type": "paragraph", "content": "text"}],
        }, skip_validation=True)
        html_doc = store.create(type="html_sandbox", name="H", data={
            "html": "<p>Hello</p>", "title": "Page",
        }, skip_validation=True)

        email_mcp = EmailDraftMCP(store)
        email = store.create(type="email_draft", name="All Types", data={
            "subject": "All", "to": ["x@y.com"], "body_html": "", "attachments": [],
        }, skip_validation=True)

        for ref_id, name in [
            (chart_id, "chart.png"),
            (table_doc.id, "data.xlsx"),
            (word_doc.id, "report.docx"),
            (html_doc.id, "page.html"),
        ]:
            result = _run(email_mcp.call_tool("email_add_attachment", {
                "document_id": email.id, "ref": ref_id, "name": name,
            }))
            assert json.loads(result)["status"] == "attachment_added"

        updated = store.get(email.id)
        assert len(updated.data["attachments"]) == 4
